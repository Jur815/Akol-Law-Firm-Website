from pathlib import Path
from datetime import date

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    BaseDocTemplate,
    Frame,
    PageBreak,
    PageTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
)


OUT_DIR = Path("docs")
BASE = "Akol_For_Legal_Services_Website_Final_Report_v1.0"
MD_PATH = OUT_DIR / f"{BASE}.md"
DOCX_PATH = OUT_DIR / f"{BASE}.docx"
PDF_PATH = OUT_DIR / f"{BASE}.pdf"

REPORT_DATE = "July 15, 2026"
DOMAIN = "https://www.akolforlegalservices.com"
FOOTER = "Prepared by Sky High Tech - Technology Made Simple"


def table(headers, rows):
    return {"type": "table", "headers": headers, "rows": rows}


sections = [
    {
        "title": "1. Cover Page",
        "blocks": [
            "Project name: Akol For Legal Services Website",
            "Client name: Akol For Legal Services",
            "Developed by: Sky High Tech",
            f"Production domain: {DOMAIN}",
            "Document title: Website Completion and Technical Handover Report",
            "Version: v1.0",
            f"Date: {REPORT_DATE}",
            "Confidentiality and intended use: This report is prepared for Akol For Legal Services and Sky High Tech for project handover, operational reference, deployment readiness, and future maintenance planning. It should not be used to disclose credentials, private tokens, or account access details.",
        ],
    },
    {
        "title": "2. Executive Summary",
        "blocks": [
            "The website provides Akol For Legal Services with a professional public presence for presenting the firm, explaining practice areas, sharing contact details, and collecting enquiries through a secure contact form.",
            "The business value is a clearer digital first impression, easier client contact, and a public channel for potential clients to understand the firm's services before initiating a legal consultation.",
            "Current completion status: the implemented codebase is production-ready from a frontend, validation, build, and serverless-contact architecture perspective. External email and domain configuration still require human setup and confirmation.",
            "Major implemented features include responsive pages, header and footer navigation, a mobile menu, clickable phone and email links, a validated contact form, a Vercel serverless contact endpoint, Resend email delivery integration, SEO metadata, robots.txt, sitemap.xml, security headers, linting, tests, and a production build.",
            "Production-readiness verdict: ready after external email configuration and final production contact-form testing.",
        ],
    },
    {
        "title": "3. Project Background",
        "blocks": [
            "The website was developed to give Akol For Legal Services a polished and accessible online presence that communicates credibility, core services, and contact options.",
            "The intended audience includes individuals, businesses, institutions, and prospective clients seeking legal advice or representation in South Sudan.",
            "The main communication objectives are to introduce the firm, explain legal service areas, provide direct contact methods, and encourage qualified enquiries.",
            "The website strengthens the firm's professional presence by presenting consistent branding, structured service information, and a secure enquiry pathway through the contact form.",
        ],
    },
    {
        "title": "4. Website Scope",
        "blocks": [
            table(
                ["Category", "Scope"],
                [
                    ["Included features", "Public informational pages, custom frontend routing, responsive layout, header/footer navigation, mobile menu, CTA links, phone and email links, contact form, serverless contact API, Resend integration, SEO files, security headers, lint/test/build scripts."],
                    ["Not included", "Online booking, payments, client portal, CMS, blog publishing system, analytics dashboard, multilingual content, legal document upload workflow, live chat, WhatsApp integration, case management, authentication, database storage."],
                    ["Current boundaries", "Static public website plus serverless email submission. Contact submissions are emailed and not stored in a database."],
                    ["Future options", "Booking, WhatsApp, blog/legal updates, analytics, CMS, testimonials, multilingual support, privacy/cookie controls, advanced anti-spam controls, and enquiry dashboard."],
                ],
            ),
        ],
    },
    {
        "title": "5. Website Structure and Pages",
        "blocks": [
            table(
                ["Route or area", "Purpose", "Main content", "Primary calls to action"],
                [
                    ["/", "Homepage and first impression.", "Firm positioning, service preview, values, contact summary, firm logo.", "Request Consultation; View Services."],
                    ["/about", "Introduce the firm and working style.", "Background, client commitment, values.", "Encourages confidence and later contact."],
                    ["/services", "Present implemented practice areas.", "Corporate and Commercial Law, Dispute Resolution, Regulatory and Compliance, Real Estate and Land, Employment and Immigration, Private Client Services.", "Start a Matter."],
                    ["/team", "Present current team and capability structure.", "Legal Counsel, Corporate Advisory Desk, Dispute Resolution Desk.", "Supports trust before contact."],
                    ["/contact", "Collect enquiries and show direct contact methods.", "Contact details, hours, validated contact form.", "Send Message."],
                    ["Fallback page", "Handle unknown routes.", "404 message and homepage link.", "Go Home."],
                    ["Header/footer", "Site-wide navigation and contact access.", "Logo, nav links, mobile menu, footer links, phone, email, address.", "Navigate to pages; call or email the firm."],
                ],
            ),
        ],
    },
    {
        "title": "6. Functional Features",
        "blocks": [
            table(
                ["Feature", "Actual implementation"],
                [
                    ["Navigation", "Custom lightweight router using browser history, popstate, and an app:navigate event. No external routing package is currently used."],
                    ["Responsive design", "CSS grid/flex layouts with media queries for desktop, tablet, and mobile."],
                    ["Mobile menu", "Header menu button toggles nav visibility and updates aria-expanded."],
                    ["Contact links", "Phone links use tel:+211923433113. Email links use mailto:info2026akollegalservices@gmail.com. WhatsApp links use https://wa.me/211912374000 with a default enquiry message."],
                    ["Calls to action", "Homepage and services CTA links route to contact or services pages."],
                    ["Contact form", "React controlled form with name, email, phone, subject, message, and hidden honeypot company field."],
                    ["Validation", "Shared frontend/server validation covers required fields, email format, phone format, message minimum length, and input length limits."],
                    ["Feedback states", "The form shows loading text, success messages, error messages, and field-level validation errors."],
                    ["Duplicate prevention", "The submit handler returns early while sending and disables form fields/button during submission."],
                    ["Spam protection", "Honeypot field plus basic in-memory IP rate limiting in the serverless function."],
                ],
            ),
        ],
    },
    {
        "title": "7. Contact Form Architecture",
        "blocks": [
            "Submission flow: Website visitor -> React contact form -> frontend validation -> POST /api/contact -> Vercel serverless function -> server-side validation -> honeypot/rate-limit checks -> Resend email API -> law firm recipient inbox.",
            "The frontend validates the form before sending and displays errors without contacting the API when required fields or formats are invalid.",
            "The API route validates the same fields server-side, rejects invalid JSON, rejects unsupported methods, checks the hidden honeypot field, and applies a basic request limit by client IP for a 15-minute window.",
            "If validation passes, the API reads RESEND_API_KEY, CONTACT_RECIPIENT_EMAIL, and CONTACT_FROM_EMAIL from server-side environment variables. It uses the visitor's email as reply_to and sends both text and safely escaped HTML email content through Resend.",
            "Success is returned only when Resend accepts the request. If Resend rejects the request, the API returns a provider failure message instead of claiming success.",
            table(
                ["Environment variable", "Purpose"],
                [
                    ["RESEND_API_KEY", "Secret API key used by the serverless function to authenticate with Resend."],
                    ["CONTACT_RECIPIENT_EMAIL", "Recipient inbox for website enquiry emails."],
                    ["CONTACT_FROM_EMAIL", "Verified Resend sender address used in outbound emails."],
                ],
            ),
            "Secrets are not stored in frontend code and are not committed to Git. .env and .env.* are ignored, while .env.example is committed as a template.",
        ],
    },
    {
        "title": "8. Technology Stack",
        "blocks": [
            table(
                ["Technology", "Actual use", "Suitability"],
                [
                    ["React 19", "Frontend UI components and stateful contact form.", "Appropriate for a modern responsive marketing and enquiry website."],
                    ["Vite 8", "Development server and production bundling.", "Fast builds and simple static deployment output."],
                    ["JavaScript", "Application, validation, and serverless API code.", "Keeps the project compact and consistent across frontend and API."],
                    ["CSS", "Custom responsive styling, layout, focus states, and brand presentation.", "Avoids unnecessary UI framework weight for a small site."],
                    ["Custom router", "Lightweight client-side routing in src/router and src/utils/router.js.", "Removes need for a routing dependency for the current route set."],
                    ["Vercel", "Intended hosting platform for frontend and serverless function.", "Matches the serverless /api/contact implementation and static frontend deployment."],
                    ["Vercel serverless functions", "api/contact.js handles secure email submission.", "Protects secrets on the server side."],
                    ["Resend", "Email delivery provider for contact submissions.", "Modern transactional email API suitable for website forms."],
                    ["ESLint", "Code quality check via npm run lint.", "Catches syntax and common JavaScript/React issues before deployment."],
                    ["Node test runner", "Tests contact validation and safe email rendering.", "No extra dependency required for current unit tests."],
                    ["Git and GitHub", "Expected source control and remote backup process.", "Recommended for deployment traceability; the inspected local folder was not recognized as a Git repository, so repository status requires manual confirmation."],
                    ["Other dev dependencies", "@vitejs/plugin-react, @rolldown/plugin-babel, Babel React Compiler preset, React type packages, React Hooks/Refresh ESLint plugins, globals.", "Support development, build, linting, and React compiler behavior."],
                ],
            ),
        ],
    },
    {
        "title": "9. System Architecture",
        "blocks": [
            "Text architecture diagram:",
            "Website visitor -> React frontend -> Vercel hosting/CDN -> Vercel serverless contact endpoint -> Resend -> Law firm email inbox",
            table(
                ["Layer", "Description"],
                [
                    ["Frontend layer", "React/Vite static site rendered in the browser with custom routing and responsive CSS."],
                    ["API layer", "Vercel serverless function at /api/contact that validates and sends messages."],
                    ["Email-delivery layer", "Resend API accepts outbound email requests from the serverless function."],
                    ["DNS and domain layer", "Production domain is www.akolforlegalservices.com. DNS values require confirmation in Bluehost and Vercel."],
                    ["Deployment layer", "Vercel build should run npm run build and serve the generated dist output plus /api serverless functions."],
                ],
            ),
        ],
    },
    {
        "title": "10. Domain and Hosting Configuration",
        "blocks": [
            table(
                ["Item", "Status"],
                [
                    ["Primary domain", "https://www.akolforlegalservices.com is used in canonical, Open Graph, robots.txt, and sitemap.xml."],
                    ["Root-domain redirect", "Requires manual confirmation in Vercel and DNS. The codebase does not prove the live redirect behavior."],
                    ["Vercel project domain", "Requires manual confirmation from the Vercel dashboard."],
                    ["Bluehost DNS management", "Referenced as required infrastructure but not verifiable from the codebase."],
                    ["Root A record", "Requires manual confirmation. Configure according to current Vercel instructions if using apex/root domain."],
                    ["www CNAME record", "Requires manual confirmation. Typically points www to the Vercel-assigned target."],
                    ["HTTPS/SSL", "Expected to be handled by Vercel once the domain is connected and verified; must be confirmed after deployment."],
                    ["DNS propagation", "May take minutes to 48 hours depending on registrar/TTL behavior."],
                ],
            ),
        ],
    },
    {
        "title": "11. SEO Implementation",
        "blocks": [
            table(
                ["SEO item", "Implementation"],
                [
                    ["Page title", "Akol For Legal Services | South Sudan Law Firm."],
                    ["Meta description", "Present in index.html and describes the legal services offering."],
                    ["Canonical URL", "https://www.akolforlegalservices.com/."],
                    ["Open Graph metadata", "og:type, og:url, og:title, og:description, and og:image are present."],
                    ["Twitter metadata", "summary card, title, and description are present."],
                    ["Favicon", "Official Akol logo favicon.ico, favicon-32x32.png, favicon-16x16.png, and apple-touch-icon.png."],
                    ["robots.txt", "Allows indexing and references the sitemap."],
                    ["sitemap.xml", "Lists /, /about, /services, /team, and /contact."],
                    ["Semantic HTML", "Uses header, nav, main, section, article, aside, footer, address, headings, labels, and buttons."],
                    ["Image alt text", "Implemented for the firm logo images."],
                    ["Indexing readiness", "Ready at code level; live indexing depends on deployment, domain resolution, and search-engine crawl timing."],
                ],
            ),
        ],
    },
    {
        "title": "12. Accessibility and User Experience",
        "blocks": [
            table(
                ["Area", "Implementation or limitation"],
                [
                    ["Responsive layouts", "Verified at mobile, tablet, laptop, and large viewport widths with no horizontal overflow in browser checks."],
                    ["Keyboard navigation", "Native links, buttons, inputs, and visible focus styles are implemented through CSS focus-visible."],
                    ["Form labels", "All visible form inputs and textarea have associated labels."],
                    ["Contrast", "Dark slate background with white/silver text and amber accents. No automated contrast audit was run, but visual contrast is intentionally high."],
                    ["Loading states", "Submit button changes to Sending... and form controls are disabled while submitting."],
                    ["Error messages", "Field-level errors and form-level alert/status messages are present."],
                    ["Mobile usability", "Mobile menu, responsive grids, and full-width mobile buttons are implemented."],
                    ["Clickable contacts", "Phone and email links are clickable in content and footer."],
                    ["Limitations", "No formal WCAG audit with assistive technology was completed. No automated axe or screen-reader test is included."],
                ],
            ),
        ],
    },
    {
        "title": "13. Security Measures",
        "blocks": [
            table(
                ["Implemented control", "Details"],
                [
                    ["Server-side environment variables", "Resend and recipient configuration are read from process.env in api/contact.js."],
                    ["No exposed API keys", "The frontend posts to /api/contact and does not contain RESEND_API_KEY."],
                    ["Input validation", "Frontend and server validate required fields, email, phone, minimum message length, and max lengths."],
                    ["Message-length limits", "Name 80, email 254, phone 32, subject 120, message 2500 characters."],
                    ["Honeypot", "Hidden company field triggers accepted no-op behavior for likely bot submissions."],
                    ["Rate limiting", "Basic in-memory IP limit of 5 requests per 15 minutes per function instance."],
                    ["Safe rendering", "HTML email content escapes user-submitted values."],
                    ["HTTP headers", "Vercel config sets nosniff, referrer policy, frame deny, permissions policy, and HSTS."],
                    ["Git ignore", ".env and .env.* are ignored; .env.example remains trackable."],
                ],
            ),
            "Recommended future improvements: durable rate limiting backed by external storage, CAPTCHA or managed bot protection if spam increases, monitoring/alerting for email failures, dependency update cadence, formal security review, and privacy/cookie notices if analytics or tracking are added.",
        ],
    },
    {
        "title": "14. Testing and Quality Assurance",
        "blocks": [
            table(
                ["Check", "Result"],
                [
                    ["npm run lint", "Passed on July 15, 2026."],
                    ["npm test", "Passed 3 tests on July 15, 2026."],
                    ["npm run build", "Passed on July 15, 2026. Production bundle generated successfully."],
                    ["Browser-width checks", "Previously verified at 390, 768, 1366, and 1920 widths with no horizontal overflow."],
                    ["Mobile-menu verification", "Previously verified menu opens and aria-expanded updates."],
                    ["Console error checks", "Previously verified no browser console errors/warnings during local checks."],
                    ["Contact API mocked acceptance", "Previously verified mocked Resend OK response returns success."],
                    ["Invalid payload rejection", "Previously verified invalid payload returns 400 with field errors."],
                ],
            ),
            "Test scope: static build, linting, validation tests, safe HTML escaping test, local browser responsiveness, menu behavior, and contact API logic with mocked email provider.",
            "Test limitations: no live production deployment test, no real Resend delivery test, no DNS/SSL verification, no formal accessibility audit, no end-to-end test suite connected to a real Vercel deployment.",
            "Recommended final production tests: submit a real contact form after Vercel environment variables are configured, confirm inbox delivery and reply-to behavior, test root-domain redirect and HTTPS, test mobile devices, and check spam/junk folder placement.",
        ],
    },
    {
        "title": "15. Deployment Process",
        "blocks": [
            "Recommended process:",
            "1. Commit the current codebase to Git.",
            "2. Push the repository to GitHub or the configured remote source.",
            "3. Connect the repository to Vercel if not already connected.",
            "4. Configure RESEND_API_KEY, CONTACT_RECIPIENT_EMAIL, and CONTACT_FROM_EMAIL in Vercel.",
            "5. Redeploy the production site after environment variables are added.",
            "6. Connect and verify the custom domain in Vercel.",
            "7. Confirm DNS records in Bluehost or the active DNS provider.",
            "8. Verify HTTPS/SSL and root-to-www redirect behavior.",
            "9. Submit a real contact form and confirm delivery to the recipient inbox.",
        ],
    },
    {
        "title": "16. Environment and Configuration Guide",
        "blocks": [
            table(
                ["Variable", "Purpose", "Example format", "Secret", "Configure in"],
                [
                    ["RESEND_API_KEY", "Authenticates serverless function with Resend.", "re_xxxxxxxxxxxxx", "Yes", "Vercel Environment Variables"],
                    ["CONTACT_RECIPIENT_EMAIL", "Inbox receiving contact submissions.", "office@example.com", "No, but operationally sensitive", "Vercel Environment Variables"],
                    ["CONTACT_FROM_EMAIL", "Verified sender address for Resend.", "Akol For Legal Services <website@example.com>", "No", "Vercel Environment Variables"],
                ],
            ),
            "Do not place real API keys in .env.example, frontend files, screenshots, documents, or Git commits.",
        ],
    },
    {
        "title": "17. Administration and Maintenance Guide",
        "blocks": [
            table(
                ["Task", "Where to update"],
                [
                    ["Update page content", "Edit the relevant files in src/pages/."],
                    ["Change phone/email/address", "Edit src/data/siteData.js."],
                    ["Update team members", "Edit the team array in src/data/siteData.js."],
                    ["Add or revise services", "Edit the services array in src/data/siteData.js."],
                    ["Update SEO metadata", "Edit index.html, public/robots.txt, and public/sitemap.xml as needed."],
                    ["Replace images", "Replace or add files in src/assets or public and update imports/links."],
                    ["Review deployments", "Use the Vercel dashboard for build logs, deployment history, and domain status."],
                    ["Rotate Resend API key", "Create a new key in Resend, update RESEND_API_KEY in Vercel, redeploy, and revoke the old key."],
                    ["Change recipient inbox", "Update CONTACT_RECIPIENT_EMAIL in Vercel and redeploy if required."],
                    ["Renew domain", "Use the registrar/DNS provider account. Bluehost was named in the requested handover scope but must be manually confirmed."],
                    ["Monitor form delivery", "Check Resend activity logs, recipient inbox, spam folder, and Vercel function logs."],
                ],
            ),
        ],
    },
    {
        "title": "18. Backup and Recovery",
        "blocks": [
            "The recommended source backup is a GitHub repository containing the website source code. The current inspected local folder was not recognized as a Git repository, so source-control status must be confirmed.",
            "Vercel deployment history can serve as an operational rollback point after the project is connected and deployed.",
            "Environment variables should be documented securely outside the repository. Do not store real secrets in this report or in Git.",
            "DNS and domain settings should be recorded by the account owner, including registrar, DNS provider, root record, www record, Resend DNS records, and renewal dates.",
            "Recovery after a failed deployment: revert to the last known good Git commit or Vercel deployment, confirm environment variables, redeploy, verify the site loads, and run the production contact-form test.",
        ],
    },
    {
        "title": "19. Known Limitations and Remaining Actions",
        "blocks": [
            table(
                ["Item", "Status"],
                [
                    ["Resend account setup", "Requires human confirmation."],
                    ["Resend sending-domain verification", "Required before using a production sender domain."],
                    ["Bluehost DNS records for Resend", "Requires manual setup/confirmation if Bluehost is the active DNS provider."],
                    ["Creation of Resend API key", "Requires human action in Resend."],
                    ["Vercel environment variables", "Must be entered in Vercel before production form delivery works."],
                    ["Real production contact-form test", "Still required after deployment and environment configuration."],
                    ["Email delivery/spam-folder confirmation", "Still required with the recipient inbox."],
                    ["Git/GitHub status", "Local folder was not recognized as a Git repository during inspection; repository status requires confirmation."],
                    ["Rate limiting", "Implemented in memory only; for high traffic, use durable storage or a managed edge/bot-protection service."],
                    ["Accessibility", "Good baseline implementation, but no formal WCAG audit or screen-reader test has been completed."],
                ],
            ),
        ],
    },
    {
        "title": "20. Future Enhancement Recommendations",
        "blocks": [
            table(
                ["Priority", "Recommendations"],
                [
                    ["Immediate", "Complete Resend verification, Vercel environment setup, production contact-form test, DNS/SSL verification, and client approval review."],
                    ["Medium-term", "Add analytics, privacy/cookie controls if tracking is introduced, CMS integration for easier content updates, testimonials, blog/legal updates, and enhanced anti-spam protection."],
                    ["Optional long-term", "Legal consultation booking, WhatsApp integration, client enquiry dashboard, multilingual support, case-study content, client portal, and richer performance monitoring."],
                ],
            ),
            "These recommendations are not currently implemented and should be scoped separately.",
        ],
    },
    {
        "title": "21. Final Readiness Assessment",
        "blocks": [
            table(
                ["Area", "Status", "Notes"],
                [
                    ["Design", "Complete", "Professional responsive visual identity is implemented."],
                    ["Responsiveness", "Complete", "Browser checks found no overflow at tested widths."],
                    ["Functionality", "Complete", "Navigation, CTAs, contact links, and form behavior are implemented."],
                    ["Contact form", "Complete pending external configuration", "Code is complete; Resend/Vercel production variables and real delivery test remain."],
                    ["SEO", "Complete", "Core metadata, sitemap, robots, favicon, and canonical domain are implemented."],
                    ["Accessibility", "Partially complete", "Baseline semantics and labels are implemented; formal WCAG audit remains."],
                    ["Security", "Complete with recommended future improvements", "Secrets server-side, validation, escaping, honeypot, rate limiting, and headers are implemented."],
                    ["Testing", "Complete for local scope", "Lint, tests, build, local browser checks, and mocked API checks passed."],
                    ["Deployment", "Complete pending external configuration", "Vercel-ready code exists; actual deployment/domain/email setup requires confirmation."],
                    ["Documentation", "Complete", "This handover report plus README and .env.example are present."],
                ],
            ),
        ],
    },
    {
        "title": "22. Final Verdict",
        "blocks": [
            "The website is ready after external email configuration.",
            "The codebase is ready for public production use on Vercel once the Resend sending domain, Resend API key, Vercel environment variables, DNS/domain setup, HTTPS, and real production contact-form delivery are confirmed.",
            "It should not be considered fully operational for client enquiries until the production contact form has been tested with the real recipient inbox.",
        ],
    },
    {
        "title": "23. Handover Checklist",
        "blocks": [
            "[ ] Code committed",
            "[ ] GitHub updated",
            "[ ] Vercel deployment complete",
            "[ ] Custom domain verified",
            "[ ] SSL active",
            "[ ] Resend domain verified",
            "[ ] Environment variables configured",
            "[ ] Contact form production-tested",
            "[ ] Recipient confirmed",
            "[ ] Client approval received",
            "[ ] Maintenance responsibility confirmed",
        ],
    },
    {
        "title": "24. Appendices",
        "blocks": [
            table(
                ["Key directory or file", "Purpose"],
                [
                    ["src/pages/", "Page-level React components."],
                    ["src/components/", "Reusable layout and common components."],
                    ["src/data/siteData.js", "Contact details, nav links, services, values, and team data."],
                    ["src/utils/contactValidation.js", "Shared contact-form validation rules."],
                    ["src/utils/router.js", "Lightweight client-side navigation helpers."],
                    ["api/contact.js", "Vercel serverless contact endpoint."],
                    ["public/robots.txt", "Search crawler instruction file."],
                    ["public/sitemap.xml", "Search-engine sitemap."],
                    ["vercel.json", "SPA rewrite and security headers."],
                    [".env.example", "Environment variable template."],
                    ["test/contactValidation.test.js", "Node tests for validation and safe email rendering."],
                ],
            ),
            table(
                ["Command", "Purpose"],
                [
                    ["npm install", "Install dependencies."],
                    ["npm run dev", "Start local development server."],
                    ["npm run lint", "Run ESLint."],
                    ["npm test", "Run Node tests."],
                    ["npm run build", "Create production build."],
                    ["npm run preview", "Preview production build locally."],
                ],
            ),
            "Environment-variable template: RESEND_API_KEY=, CONTACT_RECIPIENT_EMAIL=, CONTACT_FROM_EMAIL=.",
            "Manual test checklist: visit every page, open/close mobile menu, click phone/email links, submit empty contact form, submit valid contact form after environment setup, confirm email delivery, confirm reply-to, verify HTTPS, verify root/www domain behavior, check mobile and desktop layouts.",
            "Contact-form data flow: visitor input -> frontend validation -> /api/contact -> server validation -> spam checks -> Resend -> recipient inbox -> user-facing success or error response.",
            "Glossary: Vercel is the hosting/deployment platform; serverless function is backend code that runs on demand; Resend is the email delivery provider; environment variable is a private deployment setting; DNS connects a domain name to hosting/email services; SSL/HTTPS encrypts website traffic; sitemap helps search engines find pages.",
        ],
    },
]


def iter_toc_titles():
    return [section["title"] for section in sections if section["title"] != "1. Cover Page"]


def write_markdown():
    lines = [
        "# Akol For Legal Services Website Completion and Technical Handover Report",
        "",
        f"**Client:** Akol For Legal Services  ",
        f"**Prepared by:** Sky High Tech  ",
        f"**Production domain:** {DOMAIN}  ",
        f"**Version:** v1.0  ",
        f"**Date:** {REPORT_DATE}",
        "",
        "## Table of Contents",
        "",
    ]
    for title in iter_toc_titles():
        lines.append(f"- {title}")
    lines.append("")

    for section in sections:
        lines.append(f"## {section['title']}")
        lines.append("")
        for block in section["blocks"]:
            if isinstance(block, str):
                if block.startswith("[ ]"):
                    lines.append(f"- {block}")
                elif block and block[0].isdigit() and ". " in block[:4]:
                    lines.append(block)
                else:
                    lines.append(block)
                lines.append("")
            elif block["type"] == "table":
                headers = block["headers"]
                lines.append("| " + " | ".join(headers) + " |")
                lines.append("| " + " | ".join(["---"] * len(headers)) + " |")
                for row in block["rows"]:
                    safe = [str(cell).replace("\n", " ").replace("|", "\\|") for cell in row]
                    lines.append("| " + " | ".join(safe) + " |")
                lines.append("")
    lines.append(f"_{FOOTER}_")
    lines.append("")
    MD_PATH.write_text("\n".join(lines), encoding="utf-8")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_borders(tbl):
    tbl_pr = tbl._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "D9DEE7")


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_1)
    run._r.append(instr_text)
    run._r.append(fld_char_2)


def set_run(run, size=11, bold=False, color="000000"):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def add_docx_paragraph(doc, text, style=None, bold_label=False):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.10
    run = p.add_run(text)
    set_run(run, 11, bold_label)
    return p


def add_docx_table(doc, headers, rows):
    table_obj = doc.add_table(rows=1, cols=len(headers))
    table_obj.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_obj.autofit = False
    set_table_borders(table_obj)
    widths = [int(9360 / len(headers))] * len(headers)
    for i, header in enumerate(headers):
        cell = table_obj.rows[0].cells[i]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, "F2F4F7")
        set_cell_width(cell, widths[i])
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(header)
        set_run(run, 9.5, True, "1F4D78")
    for row in rows:
        cells = table_obj.add_row().cells
        for i, value in enumerate(row):
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_width(cells[i], widths[i])
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            run = p.add_run(str(value))
            set_run(run, 8.8)
    doc.add_paragraph()


def write_docx():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)
    for style_name, size, color in [
        ("Heading 1", 16, "2E74B5"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 12, "1F4D78"),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.text = FOOTER
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(fp.runs[0], 9, False, "666666")
    add_page_number(footer.add_paragraph())

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(120)
    title.paragraph_format.space_after = Pt(12)
    run = title.add_run("Akol For Legal Services Website")
    set_run(run, 24, True, "0B2545")
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Completion and Technical Handover Report")
    set_run(run, 16, False, "2E74B5")
    meta_lines = [
        "Client: Akol For Legal Services",
        "Developed by: Sky High Tech",
        f"Production domain: {DOMAIN}",
        "Version: v1.0",
        f"Date: {REPORT_DATE}",
        "Confidential - prepared for project handover and operational reference.",
    ]
    for line in meta_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        set_run(run, 11)
    doc.add_page_break()

    doc.add_heading("Table of Contents", level=1)
    for title_text in iter_toc_titles():
        add_docx_paragraph(doc, title_text)
    doc.add_page_break()

    for section_data in sections:
        if section_data["title"] == "1. Cover Page":
            continue
        doc.add_heading(section_data["title"], level=1)
        for block in section_data["blocks"]:
            if isinstance(block, str):
                if block.startswith("[ ]"):
                    p = doc.add_paragraph(style="List Bullet")
                    run = p.add_run(block)
                    set_run(run, 11)
                elif block and block[0].isdigit() and ". " in block[:4]:
                    p = doc.add_paragraph(style="List Number")
                    run = p.add_run(block[3:])
                    set_run(run, 11)
                else:
                    add_docx_paragraph(doc, block)
            elif block["type"] == "table":
                add_docx_table(doc, block["headers"], block["rows"])

    doc.core_properties.title = "Akol For Legal Services Website Completion and Technical Handover Report"
    doc.core_properties.author = "Sky High Tech"
    doc.save(DOCX_PATH)


def pdf_footer(canvas, doc):
    canvas.saveState()
    canvas.setFont("Helvetica", 8)
    canvas.setFillColor(colors.HexColor("#666666"))
    canvas.drawCentredString(letter[0] / 2, 0.42 * inch, FOOTER)
    canvas.drawRightString(letter[0] - inch, 0.42 * inch, f"Page {doc.page}")
    canvas.restoreState()


def ptext(text):
    return (
        str(text)
        .replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
    )


def write_pdf():
    styles = getSampleStyleSheet()
    body = ParagraphStyle(
        "Body",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=9,
        leading=12,
        spaceAfter=6,
    )
    h1 = ParagraphStyle(
        "H1",
        parent=styles["Heading1"],
        fontName="Helvetica-Bold",
        fontSize=14,
        leading=17,
        textColor=colors.HexColor("#2E74B5"),
        spaceBefore=12,
        spaceAfter=7,
    )
    title_style = ParagraphStyle(
        "Title",
        parent=styles["Title"],
        fontName="Helvetica-Bold",
        fontSize=22,
        leading=26,
        alignment=TA_CENTER,
        textColor=colors.HexColor("#0B2545"),
    )
    subtitle_style = ParagraphStyle(
        "Subtitle",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=14,
        leading=18,
        alignment=TA_CENTER,
        textColor=colors.HexColor("#2E74B5"),
        spaceAfter=12,
    )
    small = ParagraphStyle(
        "Small",
        parent=body,
        fontSize=8,
        leading=10,
    )

    doc = BaseDocTemplate(
        str(PDF_PATH),
        pagesize=letter,
        leftMargin=0.72 * inch,
        rightMargin=0.72 * inch,
        topMargin=0.78 * inch,
        bottomMargin=0.75 * inch,
    )
    frame = Frame(doc.leftMargin, doc.bottomMargin, doc.width, doc.height, id="normal")
    doc.addPageTemplates([PageTemplate(id="report", frames=[frame], onPage=pdf_footer)])
    story = []

    story.append(Spacer(1, 1.35 * inch))
    story.append(Paragraph("Akol For Legal Services Website", title_style))
    story.append(Paragraph("Completion and Technical Handover Report", subtitle_style))
    for line in [
        "Client: Akol For Legal Services",
        "Developed by: Sky High Tech",
        f"Production domain: {DOMAIN}",
        "Version: v1.0",
        f"Date: {REPORT_DATE}",
        "Confidential - prepared for project handover and operational reference.",
    ]:
        story.append(Paragraph(ptext(line), ParagraphStyle("Center", parent=body, alignment=TA_CENTER)))
    story.append(PageBreak())

    story.append(Paragraph("Table of Contents", h1))
    for title_text in iter_toc_titles():
        story.append(Paragraph(ptext(title_text), body))
    story.append(PageBreak())

    for section_data in sections:
        if section_data["title"] == "1. Cover Page":
            continue
        story.append(Paragraph(ptext(section_data["title"]), h1))
        for block in section_data["blocks"]:
            if isinstance(block, str):
                if block.startswith("[ ]"):
                    story.append(Paragraph(ptext(block), body))
                elif block and block[0].isdigit() and ". " in block[:4]:
                    story.append(Paragraph(ptext(block), body))
                else:
                    story.append(Paragraph(ptext(block), body))
            elif block["type"] == "table":
                data = [[Paragraph(ptext(h), small) for h in block["headers"]]]
                for row in block["rows"]:
                    data.append([Paragraph(ptext(cell), small) for cell in row])
                col_count = len(block["headers"])
                widths = [doc.width / col_count] * col_count
                tbl = Table(data, colWidths=widths, repeatRows=1, hAlign="CENTER", splitByRow=1)
                tbl.setStyle(
                    TableStyle(
                        [
                            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#F2F4F7")),
                            ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#1F4D78")),
                            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                            ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#D9DEE7")),
                            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                            ("LEFTPADDING", (0, 0), (-1, -1), 5),
                            ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                            ("TOPPADDING", (0, 0), (-1, -1), 4),
                            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                        ]
                    )
                )
                story.append(tbl)
                story.append(Spacer(1, 8))

    doc.build(story)


def main():
    OUT_DIR.mkdir(exist_ok=True)
    write_markdown()
    write_docx()
    write_pdf()
    print(MD_PATH)
    print(DOCX_PATH)
    print(PDF_PATH)


if __name__ == "__main__":
    main()
