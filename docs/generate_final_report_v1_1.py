from copy import deepcopy
from pathlib import Path

import generate_final_report as base_report
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    BaseDocTemplate,
    Frame,
    Image,
    PageBreak,
    PageTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
)


OUT_DIR = Path("docs")
ASSET_DIR = OUT_DIR / "report-assets" / "v1.1"
BASE = "Akol_For_Legal_Services_Website_Final_Report_v1.1_Final"
MD_PATH = OUT_DIR / f"{BASE}.md"
DOCX_PATH = OUT_DIR / f"{BASE}.docx"
PDF_PATH = OUT_DIR / f"{BASE}.pdf"

REPORT_DATE = "15 July 2026"
REPORT_DATE_LONG = "July 15, 2026"
DOMAIN = "https://www.akolforlegalservices.com"
FOOTER = "Prepared by Sky High Tech – Technology Made Simple"
VERSION = "v1.1 Final"


def table(headers, rows):
    return {"type": "table", "headers": headers, "rows": rows}


def image_block(filename, caption, description):
    return {
        "type": "image",
        "path": ASSET_DIR / filename,
        "caption": caption,
        "description": description,
    }


def title_text(section):
    return section["title"].split(". ", 1)[1]


def renumber(sections):
    numbered = []
    for index, section in enumerate(sections, start=1):
        new_section = deepcopy(section)
        new_section["title"] = f"{index}. {title_text(section)}"
        numbered.append(new_section)
    return numbered


def build_sections():
    source = deepcopy(base_report.sections)
    for section in source:
        if title_text(section) == "Cover Page":
            section["blocks"] = [
                "Project name: Akol For Legal Services Website",
                "Client name: Akol For Legal Services",
                "Developed by: Sky High Tech",
                f"Production domain: {DOMAIN}",
                "Document title: Website Completion and Technical Handover Report",
                f"Version: {VERSION}",
                f"Date: {REPORT_DATE}",
                "Confidentiality and intended use: This report is prepared for Akol For Legal Services and Sky High Tech for project handover, operational reference, deployment readiness, and future maintenance planning. It should not be used to disclose credentials, private tokens, or account access details.",
            ]
        if title_text(section) == "Executive Summary":
            section["blocks"] = [
                block.replace("external email and domain configuration still require human setup and confirmation", "external email configuration and final production contact-form testing remain")
                .replace("ready after external email configuration and final production contact-form testing", "ready after external email configuration and final production contact-form testing")
                for block in section["blocks"]
            ]
        if title_text(section) == "System Architecture":
            section["blocks"][2] = table(
                ["Layer", "Description"],
                [
                    ["Frontend layer", "React/Vite static site rendered in the browser with custom routing and responsive CSS."],
                    ["API layer", "Vercel serverless function at /api/contact that validates and sends messages."],
                    ["Email-delivery layer", "Resend API accepts outbound email requests from the serverless function."],
                    ["DNS and domain layer", "Root domain akolforlegalservices.com and primary domain www.akolforlegalservices.com have verified DNS records for Vercel. SSL certificate generation was initiated through Vercel."],
                    ["Deployment layer", "Vercel build should run npm run build and serve the generated dist output plus /api serverless functions."],
                ],
            )
        if title_text(section) == "Domain and Hosting Configuration":
            section["blocks"] = [
                table(
                    ["Item", "Verified configuration"],
                    [
                        ["Root domain", "akolforlegalservices.com."],
                        ["Primary domain", "www.akolforlegalservices.com."],
                        ["A record", "Host: @; Value: 76.76.21.21."],
                        ["CNAME record", "Host: www; Value: 869ee8e21ebb50ca.vercel-dns-017.com."],
                        ["DNS verification", "DNS verification was successfully completed."],
                        ["SSL certificate", "SSL certificate generation was initiated through Vercel."],
                        ["DNS propagation", "DNS propagation may still take time depending on registrar and network caching."],
                        ["Final production verification", "Final browser verification, HTTPS confirmation, and production contact-form testing should still be completed after deployment and environment configuration."],
                    ],
                ),
            ]
        if title_text(section) == "Known Limitations and Remaining Actions":
            section["blocks"] = [
                table(
                    ["Item", "Status"],
                    [
                        ["Resend account setup", "Requires human confirmation."],
                        ["Resend sending-domain verification", "Required before using a production sender domain."],
                        ["Bluehost DNS records for Resend", "Required if Resend domain verification records are managed through Bluehost or the active DNS provider."],
                        ["Creation of Resend API key", "Requires human action in Resend."],
                        ["Vercel environment variables", "Must be entered in Vercel before production form delivery works."],
                        ["Real production contact-form test", "Still required after deployment and environment configuration."],
                        ["Email delivery/spam-folder confirmation", "Still required with the recipient inbox."],
                        ["Git/GitHub status", "Local folder was not recognized as a Git repository during inspection; repository status requires confirmation."],
                        ["Rate limiting", "Implemented in memory only; for high traffic, use durable storage or a managed edge/bot-protection service."],
                        ["Accessibility", "Good baseline implementation, but no formal WCAG audit or screen-reader test has been completed."],
                    ],
                ),
            ]
        if title_text(section) == "Final Verdict":
            section["blocks"] = [
                "The website is ready after external email configuration.",
                "The codebase, DNS record configuration, and Vercel-oriented deployment setup are ready for public production use once the Resend sending domain, Resend API key, Vercel environment variables, SSL completion, and real production contact-form delivery are confirmed.",
                "It should not be considered fully operational for client enquiries until the production contact form has been tested with the real recipient inbox.",
            ]

    visual = {
        "title": "0. Website Visual Overview",
        "blocks": [
            image_block("homepage.png", "Figure 1. Homepage", "The homepage introduces Akol For Legal Services, presents the main value proposition, and provides primary calls to action for consultations and services."),
            image_block("about.png", "Figure 2. About Page", "The About page explains the firm's client commitment and professional values in a clear executive-facing format."),
            image_block("services.png", "Figure 3. Services Page", "The Services page lists the implemented practice areas and directs visitors toward starting a matter."),
            image_block("team.png", "Figure 4. Team Page", "The Team page presents the firm's legal capability and service desks to build trust before contact."),
            image_block("contact.png", "Figure 5. Contact Page", "The Contact page centralizes phone, email, office details, hours, and the secure enquiry form."),
            image_block("mobile-responsive.png", "Figure 6. Mobile Responsive View", "The mobile screenshot confirms that the website adapts to a small screen with accessible navigation and readable content."),
        ],
    }
    metrics = {
        "title": "0. Project Metrics and Statistics",
        "blocks": [
            table(
                ["Metric", "Actual value", "Notes"],
                [
                    ["Number of public pages", "5", "Home, About, Services, Team, Contact."],
                    ["Fallback/error page", "1", "NotFoundPage handles unknown routes."],
                    ["Reusable components", "6", "3 layout components and 3 common components in src/components/."],
                    ["API endpoints", "1", "Vercel serverless endpoint: /api/contact."],
                    ["Forms", "1", "Contact form on the Contact page."],
                    ["Environment variables", "3", "RESEND_API_KEY, CONTACT_RECIPIENT_EMAIL, CONTACT_FROM_EMAIL."],
                    ["Automated tests", "3", "Node tests in test/contactValidation.test.js."],
                    ["Implemented routes", "6", "Five public routes plus fallback route."],
                    ["Deployed domains", "2", "Root domain and primary www domain are configured."],
                    ["Build status", "Passed", "npm run build completed successfully."],
                    ["Lint status", "Passed", "npm run lint completed successfully."],
                ],
            ),
        ],
    }
    timeline = {
        "title": "0. Project Delivery Timeline",
        "blocks": [
            table(
                ["Milestone", "Status"],
                [
                    ["Requirements and Planning", "Complete"],
                    ["Design and Content Preparation", "Complete"],
                    ["Frontend Development", "Complete"],
                    ["Responsive Design Implementation", "Complete"],
                    ["Contact Form Development", "Complete"],
                    ["Security and Validation", "Complete"],
                    ["Testing and Quality Assurance", "Complete for local scope"],
                    ["Domain Configuration", "DNS verification completed; SSL certificate generation initiated through Vercel"],
                    ["Deployment Preparation", "Complete pending environment variables and final production checks"],
                    ["Production Readiness Review", "Complete; final handover version prepared"],
                ],
            ),
        ],
    }
    benefits = {
        "title": "0. Business Value and Benefits to Akol For Legal Services",
        "blocks": [
            "The website gives Akol For Legal Services a professional online presence that reflects credibility, structure, and readiness to serve clients. It helps the firm make a strong first impression before a visitor makes direct contact.",
            "The clear presentation of services, values, contact details, and firm capability improves client trust by making the firm easier to understand and easier to reach.",
            "The production domain, search-friendly metadata, sitemap, and public page structure support better visibility and discoverability as the firm builds its online presence over time.",
            "The responsive design makes the website accessible on mobile phones, tablets, laptops, and large screens, which is important for clients who browse and make enquiries from mobile devices.",
            "The contact form and clickable phone/email links create a centralized contact channel, reducing friction for prospective clients who want to begin a conversation.",
            "The website also creates a scalable digital foundation. Future services such as booking, legal updates, testimonials, analytics, WhatsApp integration, or multilingual content can be added without replacing the whole site.",
            "By presenting the firm clearly online, Akol For Legal Services gains a competitive advantage and a long-term marketing asset that can support referrals, reputation, and client acquisition.",
        ],
    }
    signature = {
        "title": "0. Signature and Approval",
        "blocks": [
            "Prepared By: Sky High Tech",
            "Founder, CEO & Lead Software Engineer: Peter Jur Makender Makech",
            "Company: Sky High Tech",
            "Slogan: Technology Made Simple",
            f"Date: {REPORT_DATE}",
            table(
                ["Approval", "Name", "Signature", "Date"],
                [
                    ["Prepared By", "Sky High Tech / Peter Jur Makender Makech", "____________________________", "15 July 2026"],
                    ["Reviewed By", "____________________________", "____________________________", "____________________________"],
                    ["Client Acceptance", "Akol For Legal Services", "____________________________", "____________________________"],
                ],
            ),
        ],
    }

    result = []
    for section in source:
        result.append(section)
        if title_text(section) == "Website Structure and Pages":
            result.extend([visual, metrics, timeline, benefits])
        if title_text(section) == "Handover Checklist":
            result.append(signature)

    return renumber(result)


sections = build_sections()


def iter_toc_titles():
    return [section["title"] for section in sections if title_text(section) != "Cover Page"]


def write_markdown():
    lines = [
        "# Akol For Legal Services Website Completion and Technical Handover Report",
        "",
        "**Client:** Akol For Legal Services  ",
        "**Prepared by:** Sky High Tech  ",
        f"**Production domain:** {DOMAIN}  ",
        f"**Version:** {VERSION}  ",
        f"**Date:** {REPORT_DATE}",
        "",
        "## Table of Contents",
        "",
    ]
    for heading in iter_toc_titles():
        lines.append(f"- {heading}")
    lines.append("")

    for section in sections:
        lines.append(f"## {section['title']}")
        lines.append("")
        for block in section["blocks"]:
            if isinstance(block, str):
                lines.append(f"- {block}" if block.startswith("[ ]") else block)
                lines.append("")
            elif block["type"] == "table":
                lines.append("| " + " | ".join(block["headers"]) + " |")
                lines.append("| " + " | ".join(["---"] * len(block["headers"])) + " |")
                for row in block["rows"]:
                    safe = [str(cell).replace("\n", " ").replace("|", "\\|") for cell in row]
                    lines.append("| " + " | ".join(safe) + " |")
                lines.append("")
            elif block["type"] == "image":
                rel_path = block["path"].as_posix()
                lines.append(f"![{block['caption']}]({rel_path})")
                lines.append("")
                lines.append(f"**{block['caption']}**  ")
                lines.append(block["description"])
                lines.append("")
    lines.append(f"_{FOOTER}_")
    lines.append("")
    MD_PATH.write_text("\n".join(lines), encoding="utf-8")


def set_run(run, size=11, bold=False, color="000000"):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_borders(tbl):
    tbl_pr = tbl._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn("w:" + edge))
        if element is None:
            element = OxmlElement("w:" + edge)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "D9DEE7")


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_1)
    run._r.append(instr_text)
    run._r.append(fld_char_2)


def add_docx_paragraph(doc, text, style=None, bold=False):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.10
    run = p.add_run(text)
    set_run(run, 11, bold)
    return p


def add_docx_table(doc, headers, rows):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    set_table_borders(tbl)
    widths = [int(9360 / len(headers))] * len(headers)
    for i, header in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, "F2F4F7")
        set_cell_width(cell, widths[i])
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(header)
        set_run(run, 9.3, True, "1F4D78")
    for row in rows:
        cells = tbl.add_row().cells
        for i, value in enumerate(row):
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_width(cells[i], widths[i])
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            run = p.add_run(str(value))
            set_run(run, 8.5)
    doc.add_paragraph()


def add_docx_image(doc, block):
    path = block["path"]
    if not path.exists():
        add_docx_paragraph(doc, f"Screenshot placeholder: {block['caption']}", bold=True)
        add_docx_paragraph(doc, "Insert screenshot before final PDF generation.")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(5.9), height=Inches(3.35))
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = caption.add_run(block["caption"])
    set_run(cap_run, 9.5, True, "1F4D78")
    add_docx_paragraph(doc, block["description"])


def write_docx():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)
    for style_name, size, color in [("Heading 1", 16, "2E74B5"), ("Heading 2", 13, "2E74B5"), ("Heading 3", 12, "1F4D78")]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.text = FOOTER
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(fp.runs[0], 9, False, "666666")
    add_page_number(footer.add_paragraph())

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(120)
    title.paragraph_format.space_after = Pt(12)
    set_run(title.add_run("Akol For Legal Services Website"), 24, True, "0B2545")
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(subtitle.add_run("Completion and Technical Handover Report"), 16, False, "2E74B5")
    for line in [
        "Client: Akol For Legal Services",
        "Developed by: Sky High Tech",
        f"Production domain: {DOMAIN}",
        f"Version: {VERSION}",
        f"Date: {REPORT_DATE}",
        "Confidential - prepared for project handover and operational reference.",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run(p.add_run(line), 11)
    doc.add_page_break()

    doc.add_heading("Table of Contents", level=1)
    for heading in iter_toc_titles():
        add_docx_paragraph(doc, heading)
    doc.add_page_break()

    for section_data in sections:
        if title_text(section_data) == "Cover Page":
            continue
        if title_text(section_data) == "Website Visual Overview":
            doc.add_page_break()
        if title_text(section_data) == "Signature and Approval":
            doc.add_page_break()
        if title_text(section_data) == "Appendices":
            doc.add_page_break()
        doc.add_heading(section_data["title"], level=1)
        for block in section_data["blocks"]:
            if isinstance(block, str):
                if block.startswith("[ ]"):
                    p = doc.add_paragraph(style="List Bullet")
                    set_run(p.add_run(block), 11)
                else:
                    add_docx_paragraph(doc, block)
            elif block["type"] == "table":
                add_docx_table(doc, block["headers"], block["rows"])
            elif block["type"] == "image":
                add_docx_image(doc, block)

    doc.core_properties.title = "Akol For Legal Services Website Completion and Technical Handover Report"
    doc.core_properties.author = "Sky High Tech"
    doc.save(DOCX_PATH)


def pdf_footer(canvas, doc):
    canvas.saveState()
    canvas.setFont("Helvetica", 8)
    canvas.setFillColor(colors.HexColor("#666666"))
    canvas.drawCentredString(letter[0] / 2, 0.42 * inch, FOOTER)
    canvas.drawRightString(letter[0] - inch, 0.42 * inch, f"Page {doc.page}")
    canvas.restoreState()


def ptext(text):
    return str(text).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def write_pdf():
    styles = getSampleStyleSheet()
    body = ParagraphStyle("Body", parent=styles["BodyText"], fontName="Helvetica", fontSize=9, leading=12, spaceAfter=6)
    h1 = ParagraphStyle("H1", parent=styles["Heading1"], fontName="Helvetica-Bold", fontSize=14, leading=17, textColor=colors.HexColor("#2E74B5"), spaceBefore=12, spaceAfter=7)
    title_style = ParagraphStyle("Title", parent=styles["Title"], fontName="Helvetica-Bold", fontSize=22, leading=26, alignment=TA_CENTER, textColor=colors.HexColor("#0B2545"))
    subtitle_style = ParagraphStyle("Subtitle", parent=styles["BodyText"], fontName="Helvetica", fontSize=14, leading=18, alignment=TA_CENTER, textColor=colors.HexColor("#2E74B5"), spaceAfter=12)
    center = ParagraphStyle("Center", parent=body, alignment=TA_CENTER)
    small = ParagraphStyle("Small", parent=body, fontSize=7.6, leading=9.2)
    cap = ParagraphStyle("Caption", parent=body, fontName="Helvetica-Bold", fontSize=8.5, leading=10.5, alignment=TA_CENTER, textColor=colors.HexColor("#1F4D78"))

    doc = BaseDocTemplate(str(PDF_PATH), pagesize=letter, leftMargin=0.72 * inch, rightMargin=0.72 * inch, topMargin=0.78 * inch, bottomMargin=0.75 * inch)
    frame = Frame(doc.leftMargin, doc.bottomMargin, doc.width, doc.height, id="normal")
    doc.addPageTemplates([PageTemplate(id="report", frames=[frame], onPage=pdf_footer)])
    story = []
    story.append(Spacer(1, 1.35 * inch))
    story.append(Paragraph("Akol For Legal Services Website", title_style))
    story.append(Paragraph("Completion and Technical Handover Report", subtitle_style))
    for line in [
        "Client: Akol For Legal Services",
        "Developed by: Sky High Tech",
        f"Production domain: {DOMAIN}",
        f"Version: {VERSION}",
        f"Date: {REPORT_DATE}",
        "Confidential - prepared for project handover and operational reference.",
    ]:
        story.append(Paragraph(ptext(line), center))
    story.append(PageBreak())
    story.append(Paragraph("Table of Contents", h1))
    for heading in iter_toc_titles():
        story.append(Paragraph(ptext(heading), body))
    story.append(PageBreak())

    for section_data in sections:
        name = title_text(section_data)
        if name == "Cover Page":
            continue
        if name in {"Website Visual Overview", "Signature and Approval"}:
            story.append(PageBreak())
        if name == "Appendices":
            story.append(PageBreak())
        story.append(Paragraph(ptext(section_data["title"]), h1))
        for block in section_data["blocks"]:
            if isinstance(block, str):
                story.append(Paragraph(ptext(block), body))
            elif block["type"] == "table":
                data = [[Paragraph(ptext(h), small) for h in block["headers"]]]
                for row in block["rows"]:
                    data.append([Paragraph(ptext(cell), small) for cell in row])
                widths = [doc.width / len(block["headers"])] * len(block["headers"])
                tbl = Table(data, colWidths=widths, repeatRows=1, hAlign="CENTER", splitByRow=1)
                tbl.setStyle(TableStyle([
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#F2F4F7")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#1F4D78")),
                    ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#D9DEE7")),
                    ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 5),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                    ("TOPPADDING", (0, 0), (-1, -1), 4),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                ]))
                story.append(tbl)
                story.append(Spacer(1, 8))
            elif block["type"] == "image":
                if block["path"].exists():
                    story.append(Image(str(block["path"]), width=doc.width, height=doc.width * 0.56))
                    story.append(Paragraph(ptext(block["caption"]), cap))
                    story.append(Paragraph(ptext(block["description"]), body))
                    story.append(Spacer(1, 7))
                else:
                    story.append(Paragraph(ptext(f"Screenshot placeholder: {block['caption']}"), cap))
                    story.append(Paragraph("Insert screenshot before final PDF generation.", body))
    doc.build(story)


def main():
    OUT_DIR.mkdir(exist_ok=True)
    write_markdown()
    write_docx()
    write_pdf()
    print(MD_PATH)
    print(DOCX_PATH)
    print(PDF_PATH)


if __name__ == "__main__":
    main()
