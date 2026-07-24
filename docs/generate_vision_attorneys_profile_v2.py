from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from generate_vision_attorneys_profile import (
    BORDER,
    GOLD,
    LIGHT_BLUE,
    LIGHT_GOLD,
    LIGHT_GRAY,
    MID_GRAY,
    NAVY,
    WHITE,
    add_callout,
    add_field,
    add_heading,
    add_key_value_table,
    add_para,
    add_service_grid,
    add_team_profile,
    add_toc,
    configure_styles,
    keep_row_together,
    paragraph_border_bottom,
    set_cell_border,
    set_cell_margins,
    set_cell_shading,
    set_document_settings,
    set_run_font,
    set_table_width,
)


OUTPUT = Path(__file__).resolve().parents[1] / "Vision_Attorneys_LF_Corporate_Profile_v2.0_Final.docx"


def set_footer(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.text = ""
    r = p.add_run("Vision Attorneys LF | Institutional Corporate Profile")
    set_run_font(r, size=8.5, color=MID_GRAY)
    spacer = p.add_run("    Page ")
    set_run_font(spacer, size=8.5, color=MID_GRAY)
    add_field(p, " PAGE ")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def add_section_divider(doc, title, subtitle):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, [6.25])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "0B2545")
    set_cell_border(cell, "0B2545", "8")
    set_cell_margins(cell, 180, 220, 180, 220)
    p = cell.paragraphs[0]
    r = p.add_run(title)
    set_run_font(r, size=15, color=WHITE, bold=True, name="Aptos Display")
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(subtitle)
    set_run_font(r2, size=9.5, color=RGBColor(236, 222, 188), italic=True)


def add_cover(doc):
    add_para(doc, "", after=8)
    band = doc.add_table(rows=1, cols=1)
    band.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(band, [6.25])
    cell = band.cell(0, 0)
    set_cell_shading(cell, "0B2545")
    set_cell_border(cell, "0B2545", "8")
    set_cell_margins(cell, 240, 240, 240, 240)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("[Insert Firm Logo]")
    set_run_font(r, size=11, color=WHITE, italic=True)
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("INSTITUTIONAL CORPORATE PROFILE")
    set_run_font(r2, size=11, color=GOLD, bold=True, name="Aptos Display")

    add_para(doc, "Vision Attorneys LF", size=31, color=NAVY, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, before=26, after=2)
    add_para(doc, "(Advocates & Commissioner for Oaths)", size=14, color=GOLD, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=10)
    rule = add_para(doc, "", after=14)
    paragraph_border_bottom(rule, "AE822A", "18", "4")
    add_para(doc, "[Insert Firm Tagline or Institutional Positioning Statement]", size=12.5, color=MID_GRAY, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=22)
    add_callout(
        doc,
        "Prepared for Institutional, Corporate, and High-Value Client Engagements",
        "A premium, fully editable profile for government agencies, NGOs, international organizations, embassies, development partners, corporate clients, investors, and high-net-worth individuals.",
    )
    add_para(doc, "", after=12)
    add_key_value_table(
        doc,
        [
            ("Head Office", "[Insert Physical Address, Juba, South Sudan]"),
            ("Telephone", "[Insert Telephone Numbers]"),
            ("Email", "[Insert Institutional Email Address]"),
            ("Website", "[Insert Website URL]"),
            ("Prepared", "[Insert Month Year]"),
        ],
        widths=(1.6, 4.55),
        header_fill=LIGHT_GOLD,
    )
    doc.add_page_break()


def add_premium_toc(doc):
    add_heading(doc, "Table of Contents", 1)
    p = doc.add_paragraph()
    add_field(p, ' TOC \\o "1-2" \\h \\z \\u ')
    for item in [
        "1. Executive Summary",
        "2. Message from the Managing Partner",
        "3. About the Firm",
        "4. Organizational Structure",
        "5. Professional Memberships & Affiliations",
        "6. Our Legal Services",
        "7. Industries Served",
        "8. Geographic Coverage",
        "9. Our Team",
        "10. Why Clients Trust Us",
        "11. Key Practice Areas",
        "12. Representative Transactions and Assignments",
        "13. Representative Clients",
        "14. Client Testimonials",
        "15. Selected Cases & Success Stories",
        "16. Firm Achievements & Milestones",
        "17. Corporate Governance & Ethics",
        "18. Corporate Social Responsibility",
        "19. Strategic Partnerships",
        "20. Strategic Vision 2030",
        "21. Frequently Asked Questions",
        "22. Contact Information",
        "23. Legal Disclaimer",
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(item)
        set_run_font(r, size=8.8, color=RGBColor(45, 53, 64))
    add_para(doc, "Note: In Microsoft Word, right-click the table above and select Update Field after editing headings or page numbers.", size=8.5, color=MID_GRAY, italic=True)
    doc.add_page_break()


def add_exec_summary(doc):
    add_heading(doc, "1. Executive Summary", 1)
    add_section_divider(doc, "Institutional Legal Profile", "A premium profile for formal engagements, proposals, tenders, and partnership discussions.")
    add_para(
        doc,
        "Vision Attorneys LF (Advocates & Commissioner for Oaths) is positioned as a professional, ethical, and institutionally capable law firm serving individuals, businesses, public institutions, non-governmental organizations, international partners, and development-focused stakeholders. The firm provides legal representation, advisory services, dispute resolution support, and documentation services with a strong commitment to integrity, confidentiality, and practical outcomes.",
    )
    add_para(doc, "[Insert firm-specific introduction, founding background, jurisdiction of practice, registration details, and institutional credentials.]")
    add_callout(
        doc,
        "Institutional Value Proposition",
        "Vision Attorneys LF combines legal discipline, client care, strong research, practical judgment, and ethical advocacy to help clients manage risk, resolve disputes, structure transactions, and protect their legal interests.",
    )


def add_managing_partner_message(doc):
    doc.add_page_break()
    add_heading(doc, "2. Message from the Managing Partner", 1)
    add_para(doc, "[Insert Managing Partner Photo]", size=9, color=MID_GRAY, italic=True, align=WD_ALIGN_PARAGRAPH.RIGHT, after=4)
    add_para(
        doc,
        "Welcome to Vision Attorneys LF. Our firm was built on a simple but powerful philosophy: legal service must be principled, practical, confidential, and responsive to the real needs of clients. Whether representing an individual, advising an institution, supporting a development partner, or guiding a corporate client through a complex matter, we approach every engagement with discipline, integrity, and respect for the rule of law.",
    )
    add_para(
        doc,
        "We believe that justice is strengthened when legal professionals combine technical competence with courage, fairness, and service. Our role is not only to interpret the law, but to help clients make informed decisions, manage risk, protect their rights, and pursue lawful solutions with confidence.",
    )
    add_para(
        doc,
        "At Vision Attorneys LF, client service is central to our identity. We value clear communication, timely updates, careful preparation, and professional confidentiality. We understand that clients often come to lawyers during important, sensitive, or high-pressure moments, and we consider it our duty to provide guidance that is both legally sound and practically useful.",
    )
    add_para(
        doc,
        "Looking ahead, our vision is to grow into a leading institutional law firm known for ethical advocacy, regional collaboration, strong legal research, and specialized legal services that support South Sudan's public, private, and development sectors. [Insert future vision specific to the firm.]",
    )
    add_para(doc, "[Insert Managing Partner Name]", size=11, color=NAVY, bold=True, before=10, after=1)
    add_para(doc, "Managing Partner, Vision Attorneys LF", size=9.5, color=GOLD, bold=True)


def add_about_and_structure(doc):
    doc.add_page_break()
    add_heading(doc, "3. About the Firm", 1)
    add_key_value_table(
        doc,
        [
            ("Firm History", "[Insert firm history, year established, founding partners, early clients, and growth milestones.]"),
            ("Vision Statement", "[Insert vision statement focused on justice, institutional excellence, and trusted legal service.]"),
            ("Mission Statement", "[Insert mission statement describing legal representation, advisory support, and client-centered solutions.]"),
            ("Core Values", "Integrity; confidentiality; professionalism; accountability; service; excellence; respect for the rule of law."),
            ("Areas of Specialization", "Commercial law, litigation, property, employment, family law, legal advisory, mediation, arbitration, and Commissioner for Oaths services."),
        ],
        widths=(1.65, 4.55),
        header_fill=LIGHT_BLUE,
    )
    add_heading(doc, "4. Organizational Structure", 1)
    add_para(doc, "The organizational structure below may be customized to reflect the firm's actual staffing, reporting lines, and operational departments.")
    add_org_chart(doc)


def add_org_chart(doc):
    levels = [
        ("Managing Partner", "Strategic leadership, client relations, institutional oversight"),
        ("Senior Advocates", "Complex matters, litigation leadership, supervision"),
        ("Associate Advocates", "Client matters, research, drafting, appearances"),
        ("Legal Researchers", "Case law, statutory review, legal opinions"),
        ("Legal Assistants", "Documentation, filing, client coordination"),
        ("Administrative Staff", "Office management, finance, scheduling, records"),
    ]
    table = doc.add_table(rows=len(levels), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, [2.25, 3.9])
    for i, (role, purpose) in enumerate(levels):
        c0, c1 = table.rows[i].cells
        keep_row_together(table.rows[i])
        for cell in (c0, c1):
            set_cell_border(cell, "D7DEE8", "4")
            set_cell_margins(cell, 70, 130, 70, 130)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_shading(c0, "0B2545" if i == 0 else ("F5EFE3" if i % 2 else "EEF3F8"))
        r0 = c0.paragraphs[0].add_run(role)
        set_run_font(r0, size=9.4, color=WHITE if i == 0 else NAVY, bold=True)
        r1 = c1.paragraphs[0].add_run(purpose)
        set_run_font(r1, size=8.7, color=RGBColor(45, 53, 64))


def add_memberships_services_industries_geo(doc):
    doc.add_page_break()
    add_heading(doc, "5. Professional Memberships & Affiliations", 1)
    add_key_value_table(
        doc,
        [
            ("South Sudan Bar Association", "[Insert membership status, registration number, and active standing details.]"),
            ("East Africa Law Society", "[Insert membership or affiliation details where applicable.]"),
            ("International Bar Association", "[Insert membership, participation, or future affiliation status where applicable.]"),
            ("Other Legal Bodies", "[Insert courts, professional associations, legal networks, chambers, or institutional affiliations.]"),
        ],
        widths=(2.15, 4.05),
        header_fill=LIGHT_GOLD,
    )
    add_heading(doc, "6. Our Legal Services", 1)
    services = [
        ("CL", "Corporate and Commercial Law", "Company formation, governance, compliance, commercial contracts, shareholder matters, restructurings, and transaction support."),
        ("LT", "Civil Litigation", "Claims, defenses, pleadings, court representation, injunctions, enforcement, settlement strategy, and dispute management."),
        ("CD", "Criminal Defense", "Advisory and representation for suspects, accused persons, bail, trial preparation, appeals, and rights protection."),
        ("FL", "Family Law", "Divorce, custody, maintenance, succession, adoption, family settlements, and confidential family advisory support."),
        ("EL", "Employment and Labour Law", "Employment contracts, HR policies, disciplinary procedures, termination, labour disputes, and compliance advisory."),
        ("PL", "Property and Land Law", "Land due diligence, leases, transfers, title review, property disputes, and real estate transaction support."),
        ("CR", "Contract Drafting and Review", "Drafting, review, negotiation, risk allocation, enforceability checks, and plain-language contract advice."),
        ("AM", "Arbitration and Mediation", "Alternative dispute resolution, mediation strategy, arbitration preparation, negotiation, and settlement support."),
        ("LA", "Legal Advisory Services", "Legal opinions, compliance guidance, institutional retainers, regulatory advice, and legal risk management."),
        ("CO", "Commissioner for Oaths Services", "Affidavits, declarations, oaths, witnessing, certifications, and documentation support as permitted by law."),
        ("AS", "Additional Services", "[Insert additional approved legal services, sector-specific capabilities, and emerging practice areas.]"),
    ]
    add_service_grid(doc, services)
    add_heading(doc, "7. Industries Served", 1)
    add_industry_grid(doc)
    add_heading(doc, "8. Geographic Coverage", 1)
    add_key_value_table(
        doc,
        [
            ("Juba", "Head office support, institutional meetings, court representation, client advisory, and documentation services."),
            ("Central Equatoria", "Regional advisory, property matters, dispute support, institutional engagement, and client representation."),
            ("Western Equatoria", "Representation and advisory support through case-specific arrangements and partner coordination."),
            ("Eastern Equatoria", "Legal support for client matters requiring regional coordination and field-level engagement."),
            ("Nationwide Representation", "Representation across South Sudan subject to jurisdiction, logistics, and formal engagement terms."),
            ("Regional / East African Legal Support", "Cross-border support through professional networks, referrals, and regional collaboration where appropriate."),
        ],
        widths=(2.0, 4.2),
        header_fill=LIGHT_BLUE,
    )


def add_industry_grid(doc):
    items = [
        ("NGOs and Humanitarian Organizations", "Compliance, contracts, employment, partnerships, registrations, and operational legal support."),
        ("Government Institutions", "Advisory, documentation, dispute support, procurement-related legal review, and institutional representation."),
        ("Financial Institutions", "Security documentation, recovery matters, regulatory compliance, contracts, and dispute resolution."),
        ("Construction and Engineering Companies", "Project contracts, claims, land issues, employment matters, and commercial disputes."),
        ("Telecommunications", "Commercial agreements, compliance, employment, licensing support, and dispute management."),
        ("Oil and Gas", "Regulatory advisory, contract review, land and community matters, employment, and compliance support."),
        ("Agriculture", "Land, supply contracts, cooperatives, employment, regulatory issues, and transaction support."),
        ("Real Estate and Property Development", "Due diligence, leases, sales, transfers, title review, and property disputes."),
        ("Education Institutions", "Governance, employment, policies, contracts, compliance, and dispute resolution."),
        ("SMEs and Startups", "Formation, contracts, advisory retainers, employment, intellectual property referrals, and growth-stage legal support."),
    ]
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, [3.08, 3.08])
    for idx in range(0, len(items), 2):
        row = table.add_row()
        keep_row_together(row)
        for col in range(2):
            cell = row.cells[col]
            set_cell_border(cell, "D7DEE8", "4")
            set_cell_margins(cell, 120, 145, 120, 145)
            set_cell_shading(cell, "FFFFFF" if (idx + col) % 2 == 0 else "FBFCFE")
            if idx + col < len(items):
                title, body = items[idx + col]
                r = cell.paragraphs[0].add_run(title)
                set_run_font(r, size=9.7, color=NAVY, bold=True)
                p = cell.add_paragraph()
                p.paragraph_format.space_after = Pt(0)
                rb = p.add_run(body)
                set_run_font(rb, size=8.8, color=RGBColor(45, 53, 64))


def add_team_trust_practice(doc):
    doc.add_page_break()
    add_heading(doc, "9. Our Team", 1)
    add_para(doc, "The firm may replace the placeholders below with approved biographies, professional portraits, qualifications, bar admissions, and selected experience.")
    add_team_profile(doc, "Managing Partner Profile Template", "[Insert Managing Partner Name]", "Managing Partner")
    add_team_profile(doc, "Senior Advocate Profile Template", "[Insert Senior Advocate Name]", "Senior Advocate")
    add_team_profile(doc, "Associate Advocate Profile Template", "[Insert Associate Advocate Name]", "Associate Advocate")
    add_team_profile(doc, "Legal Researcher / Legal Assistant Template", "[Insert Team Member Name]", "[Insert Position]")
    add_heading(doc, "10. Why Clients Trust Us", 1)
    add_trust_grid(doc)
    add_heading(doc, "11. Key Practice Areas", 1)
    for heading, body in [
        ("11.1 Commercial Law", "Commercial advisory for companies, institutions, entrepreneurs, and investors, including contracts, governance, compliance, and transaction support."),
        ("11.2 Litigation", "Representation in civil, commercial, employment, land, family, and institutional disputes through disciplined case preparation and strategic advocacy."),
        ("11.3 Property Law", "Support for land due diligence, leases, transfers, title review, property development, disputes, and regulatory issues."),
        ("11.4 Employment Law", "Guidance for employers and employees on employment contracts, workplace policies, termination, disciplinary process, and labour disputes."),
        ("11.5 Family Law", "Confidential legal support for family matters, including divorce, custody, maintenance, succession, adoption, and negotiated family settlements."),
    ]:
        add_heading(doc, heading, 2)
        add_para(doc, body + " [Insert firm-specific examples, credentials, and jurisdictional details.]")


def add_trust_grid(doc):
    items = [
        ("CONF", "Confidentiality", "Strict handling of sensitive legal, personal, commercial, and institutional information."),
        ("INT", "Integrity", "Ethical advocacy, honest advice, and respect for professional responsibility."),
        ("TIME", "Timely Service", "Prompt communication, clear next steps, and organized matter management."),
        ("RES", "Strong Legal Research", "Careful review of laws, regulations, case law, and factual records."),
        ("TEAM", "Experienced Team", "Collaborative legal professionals supporting complex and routine matters."),
        ("PRAC", "Practical Solutions", "Advice that balances legal rights, risk, cost, timing, and client objectives."),
        ("CARE", "Client-Centered Approach", "Responsive service built around client priorities and informed decision-making."),
    ]
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, [3.08, 3.08])
    for idx in range(0, len(items), 2):
        row = table.add_row()
        keep_row_together(row)
        for col in range(2):
            cell = row.cells[col]
            set_cell_border(cell, "D7DEE8", "4")
            set_cell_margins(cell, 120, 145, 120, 145)
            set_cell_shading(cell, LIGHT_GOLD if (idx + col) % 2 == 0 else LIGHT_BLUE)
            if idx + col < len(items):
                icon, title, body = items[idx + col]
                r = cell.paragraphs[0].add_run(f"{icon}  {title}")
                set_run_font(r, size=9.8, color=NAVY, bold=True)
                p = cell.add_paragraph()
                p.paragraph_format.space_after = Pt(0)
                rb = p.add_run(body)
                set_run_font(rb, size=8.8, color=RGBColor(45, 53, 64))


def add_transactions_clients_milestones(doc):
    doc.add_page_break()
    add_heading(doc, "12. Representative Transactions and Assignments", 1)
    add_key_value_table(
        doc,
        [
            ("Corporate Advisory", "[Insert company advisory, governance, registration, compliance, or board-support assignments.]"),
            ("Contract Negotiations", "[Insert major contract negotiation support, supplier agreements, service contracts, or partnership agreements.]"),
            ("Due Diligence", "[Insert land, corporate, employment, regulatory, or transaction due diligence assignments.]"),
            ("Regulatory Compliance", "[Insert licensing, statutory filings, policy review, and compliance advisory work.]"),
            ("Land Transactions", "[Insert leases, purchases, transfers, title review, or property development matters.]"),
            ("Employment Matters", "[Insert employment advisory, termination, workplace investigation, dispute, or policy work.]"),
            ("Litigation and Dispute Resolution", "[Insert representative disputes, arbitration, mediation, settlement, or court matters.]"),
        ],
        widths=(2.2, 4.0),
        header_fill=LIGHT_BLUE,
    )
    add_heading(doc, "13. Representative Clients", 1)
    add_para(doc, "Client names should be included only with authorization. The categories below may be used where confidentiality or professional responsibility requires anonymized presentation.")
    add_key_value_table(
        doc,
        [
            ("Government Institutions", "[Insert names or anonymized categories subject to permission.]"),
            ("NGOs", "[Insert humanitarian, development, or civil society sectors served.]"),
            ("Private Companies", "[Insert industries or approved company names.]"),
            ("Individuals", "[Insert high-net-worth individuals, entrepreneurs, families, employees, or landowners.]"),
            ("International Organizations", "[Insert international organizations, embassies, agencies, or development partners.]"),
        ],
        widths=(2.0, 4.2),
        header_fill=LIGHT_GOLD,
    )
    add_heading(doc, "14. Client Testimonials", 1)
    add_testimonials(doc)
    add_heading(doc, "15. Selected Cases & Success Stories", 1)
    add_callout(doc, "Confidentiality Disclaimer", "Case information may be presented subject to client confidentiality obligations.")
    add_key_value_table(
        doc,
        [
            ("Matter Title", "[Insert anonymized or approved matter title.]"),
            ("Client Sector", "[Insert sector or client category.]"),
            ("Legal Issue", "[Insert concise description of legal challenge.]"),
            ("Firm Role", "[Insert advisory, litigation, negotiation, drafting, or compliance role.]"),
            ("Outcome / Value", "[Insert result, settlement, risk avoided, commercial value, or institutional benefit.]"),
        ],
        widths=(1.85, 4.35),
        header_fill=LIGHT_GRAY,
    )
    add_heading(doc, "16. Firm Achievements & Milestones", 1)
    add_timeline(doc)


def add_testimonials(doc):
    testimonials = [
        ("Institutional Client", "[Insert testimonial from government agency, NGO, embassy, development partner, or institutional client.]"),
        ("Corporate Client", "[Insert testimonial from company, investor, SME, or commercial client.]"),
        ("Individual Client", "[Insert testimonial from high-net-worth individual, family client, property client, or private client.]"),
        ("Partner / Referrer", "[Insert testimonial from strategic partner, referral partner, or professional collaborator.]"),
    ]
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, [3.08, 3.08])
    for idx in range(0, len(testimonials), 2):
        row = table.add_row()
        keep_row_together(row)
        for col in range(2):
            cell = row.cells[col]
            set_cell_border(cell, "D7DEE8", "4")
            set_cell_margins(cell, 115, 145, 115, 145)
            set_cell_shading(cell, LIGHT_GOLD if (idx + col) % 2 == 0 else LIGHT_BLUE)
            if idx + col < len(testimonials):
                role, quote = testimonials[idx + col]
                r = cell.paragraphs[0].add_run(role)
                set_run_font(r, size=9.5, color=NAVY, bold=True)
                p = cell.add_paragraph()
                p.paragraph_format.space_after = Pt(0)
                rq = p.add_run(quote)
                set_run_font(rq, size=8.8, color=RGBColor(45, 53, 64), italic=True)


def add_timeline(doc):
    milestones = [
        ("Year Established", "[Insert year established and founding context.]"),
        ("Major Growth Milestones", "[Insert expansion of team, office capacity, client portfolio, or practice systems.]"),
        ("Landmark Cases", "[Insert approved landmark cases or anonymized significant matters.]"),
        ("Expansion of Services", "[Insert new practice areas, Commissioner for Oaths services, ADR, or institutional advisory services.]"),
        ("Strategic Partnerships", "[Insert referral networks, professional partnerships, or institutional collaborations.]"),
    ]
    table = doc.add_table(rows=len(milestones), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, [1.9, 4.3])
    for i, (label, detail) in enumerate(milestones):
        c0, c1 = table.rows[i].cells
        keep_row_together(table.rows[i])
        for cell in (c0, c1):
            set_cell_border(cell, "D7DEE8", "4")
            set_cell_margins(cell, 130, 150, 130, 150)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_shading(c0, "0B2545")
        r0 = c0.paragraphs[0].add_run(label)
        set_run_font(r0, size=9.5, color=WHITE, bold=True)
        r1 = c1.paragraphs[0].add_run(detail)
        set_run_font(r1, size=9.5, color=RGBColor(45, 53, 64))


def add_governance_csr_vision(doc):
    doc.add_page_break()
    add_heading(doc, "17. Corporate Governance & Ethics", 1)
    add_para(doc, "Vision Attorneys LF is committed to professional conduct, ethical representation, legal compliance, and responsible client service. This section may be customized to reflect the firm's formal policies and internal procedures.")
    add_key_value_table(
        doc,
        [
            ("Professional Conduct", "The firm expects all personnel to observe professional discipline, respect court processes, and uphold the dignity of legal practice."),
            ("Anti-Corruption Commitment", "The firm does not support bribery, improper influence, conflict-driven representation, or unlawful shortcuts."),
            ("Client Confidentiality", "Client information, documents, instructions, and legal strategy are handled with discretion and professional care."),
            ("Conflict-of-Interest Policy", "Potential conflicts should be checked before accepting instructions, and clients should be informed where applicable."),
            ("Legal Compliance Standards", "The firm works within applicable laws, professional rules, procedural requirements, and engagement agreements."),
        ],
        widths=(2.0, 4.2),
        header_fill=LIGHT_BLUE,
    )
    add_heading(doc, "18. Corporate Social Responsibility", 1)
    add_key_value_table(
        doc,
        [
            ("Pro Bono Services", "[Insert pro bono policy, eligible beneficiaries, legal aid priorities, and public-interest work.]"),
            ("Legal Awareness Campaigns", "[Insert workshops, radio programs, community briefings, or legal education activities.]"),
            ("Community Outreach", "[Insert outreach programs for vulnerable groups, local communities, or civil society partners.]"),
            ("Youth Mentorship", "[Insert mentorship, internships, law student support, career talks, and professional development initiatives.]"),
            ("Access to Justice Initiatives", "[Insert programs improving access to legal information, representation, or referral pathways.]"),
        ],
        widths=(2.0, 4.2),
        header_fill=LIGHT_GOLD,
    )
    add_heading(doc, "19. Strategic Partnerships", 1)
    add_para(doc, "[Insert local and international partners, referral networks, development partners, professional associations, technical collaborators, and institutional relationships.]")
    add_heading(doc, "20. Strategic Vision 2030", 1)
    add_key_value_table(
        doc,
        [
            ("Expansion Plans", "[Insert planned office growth, practice expansion, regional reach, or institutional client strategy.]"),
            ("Technology Adoption", "[Insert plans for digital case management, legal research tools, secure document management, and client communication systems.]"),
            ("Regional Partnerships", "[Insert East African legal networks, cross-border collaboration, and referral partnerships.]"),
            ("Specialized Legal Services", "[Insert planned specialization in sectors such as energy, infrastructure, humanitarian operations, corporate advisory, or investment law.]"),
            ("Institutional Growth", "[Insert plans for governance systems, training, quality assurance, and long-term professional development.]"),
        ],
        widths=(2.0, 4.2),
        header_fill=LIGHT_BLUE,
    )


def add_faq_contact_disclaimer(doc):
    doc.add_page_break()
    add_heading(doc, "21. Frequently Asked Questions", 1)
    add_key_value_table(
        doc,
        [
            ("How do clients engage the firm?", "Clients may contact the firm through the official telephone number, email address, physical office, or referral. A formal engagement letter may be required before work begins."),
            ("What is the consultation process?", "The firm conducts an initial review of the issue, confirms whether it can assist, checks conflicts, discusses scope and fees, and agrees on next steps."),
            ("What is the confidentiality policy?", "Information shared with the firm is handled confidentially, subject to applicable law, professional obligations, and formal engagement terms."),
            ("How does legal representation work?", "Representation may include fact review, legal research, strategy development, drafting, filing, negotiation, hearings, settlement discussions, and client updates."),
        ],
        widths=(2.2, 4.0),
        header_fill=LIGHT_BLUE,
    )
    add_heading(doc, "22. Contact Information", 1)
    add_key_value_table(
        doc,
        [
            ("Physical Address", "[Insert Physical Address]"),
            ("Postal Address", "[Insert Postal Address]"),
            ("Telephone Numbers", "[Insert Telephone Numbers]"),
            ("Email Address", "[Insert Email Address]"),
            ("Website", "[Insert Website]"),
            ("Facebook", "[Insert Facebook Page URL]"),
            ("LinkedIn", "[Insert LinkedIn Page URL]"),
            ("Working Hours", "[Insert Working Hours]"),
        ],
        widths=(1.85, 4.35),
        header_fill=LIGHT_GOLD,
    )
    doc.add_page_break()
    add_heading(doc, "23. Legal Disclaimer", 1)
    add_callout(
        doc,
        "Important Notice",
        "This profile is intended for informational purposes only and does not constitute legal advice. Engagements are governed by formal agreements and applicable laws.",
    )
    add_para(doc, "No reader should rely on this profile as a substitute for legal advice from a qualified advocate after review of the specific facts, applicable law, and formal engagement terms.")
    add_para(doc, "Information about clients, cases, transactions, and assignments should be published only where permitted by professional responsibility, client authorization, confidentiality obligations, and applicable law.")
    add_para(doc, "", after=24)
    add_section_divider(doc, "Vision Attorneys LF", "[Insert final institutional closing statement, motto, or call to action.]")


def build_doc():
    doc = Document()
    configure_styles(doc)
    section = set_document_settings(doc)
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.62)
    section.left_margin = Inches(0.82)
    section.right_margin = Inches(0.82)
    set_footer(section)
    add_cover(doc)
    add_premium_toc(doc)
    add_exec_summary(doc)
    add_managing_partner_message(doc)
    add_about_and_structure(doc)
    add_memberships_services_industries_geo(doc)
    add_team_trust_practice(doc)
    add_transactions_clients_milestones(doc)
    add_governance_csr_vision(doc)
    add_faq_contact_disclaimer(doc)
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build_doc())
