from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path(__file__).resolve().parents[1] / "Vision_Attorneys_LF_Corporate_Profile_Template.docx"

NAVY = RGBColor(11, 37, 69)
GOLD = RGBColor(174, 130, 42)
LIGHT_GOLD = "F5EFE3"
LIGHT_BLUE = "EEF3F8"
LIGHT_GRAY = "F4F6F8"
MID_GRAY = RGBColor(89, 96, 105)
WHITE = RGBColor(255, 255, 255)
BORDER = "CBD5E1"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=BORDER, size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:{}".format(edge)
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=100, start=140, bottom=100, end=140):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def keep_row_together(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def set_run_font(run, size=None, color=None, bold=None, italic=None, name="Aptos"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_field(paragraph, instr):
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instr
    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_separate)
    run._r.append(fld_char_end)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run_font(run, size=9, color=MID_GRAY)
    add_field(paragraph, " PAGE ")


def paragraph_border_bottom(paragraph, color="AE822A", size="10", space="4"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)


def add_para(doc, text="", style=None, size=None, color=None, bold=None, italic=None, align=None, before=None, after=None):
    p = doc.add_paragraph(style=style)
    if text:
        run = p.add_run(text)
        set_run_font(run, size=size, color=color, bold=bold, italic=italic)
    if align is not None:
        p.alignment = align
    if before is not None:
        p.paragraph_format.space_before = Pt(before)
    if after is not None:
        p.paragraph_format.space_after = Pt(after)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    return p


def add_placeholder_box(doc, label, height_rows=3):
    table = doc.add_table(rows=height_rows, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, [6.1])
    for row in table.rows:
        cell = row.cells[0]
        set_cell_shading(cell, "FAFAFA")
        set_cell_border(cell, "B8C2CC", "8")
        set_cell_margins(cell, 140, 180, 140, 180)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cell = table.cell(0, 0)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(label)
    set_run_font(r, size=10, color=MID_GRAY, italic=True)
    return table


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [6.25])
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_BLUE)
    set_cell_border(cell, "D8E1EA", "6")
    set_cell_margins(cell, 140, 180, 140, 180)
    p = cell.paragraphs[0]
    r = p.add_run(title)
    set_run_font(r, size=10.5, color=NAVY, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    set_run_font(r2, size=10, color=RGBColor(42, 50, 60))
    return table


def add_key_value_table(doc, rows, widths=(1.85, 4.35), header_fill=None):
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, list(widths))
    for i, (label, value) in enumerate(rows):
        c0, c1 = table.rows[i].cells
        for cell in (c0, c1):
            set_cell_border(cell, BORDER, "4")
            set_cell_margins(cell, 100, 140, 100, 140)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_shading(c0, header_fill or LIGHT_GRAY)
        c0.paragraphs[0].text = ""
        r0 = c0.paragraphs[0].add_run(label)
        set_run_font(r0, size=9.5, color=NAVY, bold=True)
        c1.paragraphs[0].text = ""
        r1 = c1.paragraphs[0].add_run(value)
        set_run_font(r1, size=9.5, color=RGBColor(35, 42, 52))
    return table


def add_bullet_list(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(f" {item}")
        set_run_font(r, size=10.5, color=RGBColor(35, 42, 52))


def add_service_grid(doc, services):
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, [3.08, 3.08])
    for idx in range(0, len(services), 2):
        row = table.add_row()
        for col in range(2):
            cell = row.cells[col]
            set_cell_border(cell, "D7DEE8", "4")
            set_cell_margins(cell, 130, 150, 130, 150)
            set_cell_shading(cell, "FFFFFF" if (idx + col) % 2 == 0 else "FBFCFE")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            if idx + col < len(services):
                icon, title, body = services[idx + col]
                p = cell.paragraphs[0]
                r = p.add_run(f"{icon}  {title}")
                set_run_font(r, size=10.5, color=NAVY, bold=True)
                p2 = cell.add_paragraph()
                p2.paragraph_format.space_after = Pt(0)
                r2 = p2.add_run(body)
                set_run_font(r2, size=9.2, color=RGBColor(45, 53, 64))
            else:
                cell.text = ""
    return table


def add_team_profile(doc, title, name_placeholder, role_placeholder):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, [1.55, 4.65])
    keep_row_together(table.rows[0])
    photo, info = table.rows[0].cells
    for cell in (photo, info):
        set_cell_border(cell, "D7DEE8", "4")
        set_cell_margins(cell, 120, 140, 120, 140)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    set_cell_shading(photo, "FAFAFA")
    p = photo.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("[Insert Professional Photo]")
    set_run_font(r, size=8.8, color=MID_GRAY, italic=True)
    p = info.paragraphs[0]
    r = p.add_run(title)
    set_run_font(r, size=10.5, color=NAVY, bold=True)
    for label, value in [
        ("Name", name_placeholder),
        ("Position", role_placeholder),
        ("Qualifications", "[Insert degrees, bar admission, certifications, and professional memberships]"),
        ("Experience", "[Insert years of practice, notable legal experience, and sector exposure]"),
        ("Expertise", "[Insert principal practice areas and specialist capabilities]"),
    ]:
        p2 = info.add_paragraph()
        p2.paragraph_format.space_after = Pt(1)
        r1 = p2.add_run(f"{label}: ")
        set_run_font(r1, size=9.2, color=NAVY, bold=True)
        r2 = p2.add_run(value)
        set_run_font(r2, size=9.2, color=RGBColor(45, 53, 64))
    return table


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor(35, 42, 52)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for level, size, before, after, color in [
        (1, 16, 15, 7, NAVY),
        (2, 13, 10, 5, NAVY),
        (3, 11.5, 6, 3, RGBColor(69, 82, 98)),
    ]:
        style = styles[f"Heading {level}"]
        style.font.name = "Aptos Display"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Aptos Display")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos Display")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Aptos"
        style.font.size = Pt(10.2)
        style.paragraph_format.left_indent = Inches(0.35)
        style.paragraph_format.first_line_indent = Inches(-0.17)
        style.paragraph_format.space_after = Pt(3)


def set_document_settings(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.32)
    section.different_first_page_header_footer = True
    return section


def set_footer(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.text = ""
    left = p.add_run("Vision Attorneys LF Corporate Profile Template")
    set_run_font(left, size=8.5, color=MID_GRAY)
    tab = p.add_run("\t")
    set_run_font(tab, size=8.5, color=MID_GRAY)
    add_page_number(p)


def add_cover(doc):
    add_para(doc, "", after=6)
    logo = doc.add_table(rows=1, cols=1)
    logo.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(logo, [1.65])
    cell = logo.cell(0, 0)
    set_cell_border(cell, "AE822A", "10")
    set_cell_margins(cell, 260, 160, 260, 160)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("[Insert Firm Logo]")
    set_run_font(r, size=10, color=MID_GRAY, italic=True)

    add_para(doc, "Vision Attorneys LF", size=29, color=NAVY, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, before=18, after=2)
    add_para(doc, "(Advocates & Commissioner for Oaths)", size=14, color=GOLD, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=12)
    rule = add_para(doc, "", after=12)
    paragraph_border_bottom(rule, "AE822A", "16", "4")
    add_para(doc, "[Insert Tagline or Firm Motto]", size=12.5, color=MID_GRAY, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=22)

    add_callout(
        doc,
        "Corporate Profile Template",
        "A fully editable sample profile prepared for client presentations, business development, proposals, website content, and institutional introductions.",
    )
    add_para(doc, "", after=10)
    contact = [
        ("Physical Address", "[Insert Office Location]"),
        ("Telephone", "[Insert Telephone Numbers]"),
        ("Email", "[Insert Email Address]"),
        ("Website", "[Insert Website URL]"),
    ]
    add_key_value_table(doc, contact, widths=(1.75, 4.1), header_fill=LIGHT_GOLD)
    add_para(doc, "", after=12)
    add_para(doc, "[Insert City, Country]  |  [Insert Month Year]", size=9.5, color=MID_GRAY, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()


def add_toc(doc):
    add_heading(doc, "Table of Contents", 1)
    p = doc.add_paragraph()
    add_field(p, ' TOC \\o "1-2" \\h \\z \\u ')
    fallback = [
        "1. Executive Summary",
        "2. About the Firm",
        "3. Our Legal Services",
        "4. Our Team",
        "5. Why Choose Vision Attorneys LF",
        "6. Key Practice Areas",
        "7. Representative Clients",
        "8. Selected Cases & Success Stories",
        "9. Corporate Social Responsibility (CSR)",
        "10. Strategic Partnerships",
        "11. Firm Capacity and Competitive Advantage",
        "12. Contact Information",
    ]
    for item in fallback:
        line = doc.add_paragraph()
        line.paragraph_format.space_after = Pt(1)
        r = line.add_run(item)
        set_run_font(r, size=9.4, color=RGBColor(45, 53, 64))
    add_para(doc, "Note: In Microsoft Word, right-click the table above and select Update Field after editing headings or page numbers.", size=9, color=MID_GRAY, italic=True)
    doc.add_page_break()


def add_executive_summary(doc):
    add_heading(doc, "1. Executive Summary", 1)
    add_callout(
        doc,
        "Profile Purpose",
        "This template presents Vision Attorneys LF as a modern, ethical, and client-focused legal practice. All bracketed content is editable and should be replaced with firm-specific information before external publication.",
    )
    add_para(
        doc,
        "Vision Attorneys LF (Advocates & Commissioner for Oaths) is a professional legal practice committed to delivering practical, timely, and confidential legal solutions to individuals, businesses, institutions, and community organizations. The firm combines sound legal analysis with responsive client service and a clear understanding of the commercial, personal, and public-interest contexts in which legal matters arise.",
    )
    add_para(
        doc,
        "[Insert a concise introduction to the firm, including its founding background, jurisdiction of practice, principal office location, and the types of clients served.]"
    )
    add_para(
        doc,
        "The firm's mission is to provide dependable legal representation, strategic advisory support, and accessible legal guidance while upholding the highest standards of integrity, professionalism, and respect for client confidentiality."
    )


def add_about(doc):
    add_heading(doc, "2. About the Firm", 1)
    add_heading(doc, "2.1 Firm History", 2)
    add_para(doc, "[Insert the founding history of Vision Attorneys LF, including the year of establishment, founding partners, early areas of practice, and major growth milestones.]")
    add_heading(doc, "2.2 Vision Statement", 2)
    add_para(doc, "[Insert Vision Statement: e.g., To be a trusted legal partner known for excellence, ethical advocacy, and practical solutions that advance justice and protect client interests.]")
    add_heading(doc, "2.3 Mission Statement", 2)
    add_para(doc, "[Insert Mission Statement: e.g., To provide high-quality legal services through professional counsel, diligent representation, and client-centered strategies.]")
    add_heading(doc, "2.4 Core Values", 2)
    add_bullet_list(doc, [
        "Integrity: We act honestly, ethically, and transparently in every client engagement.",
        "Professional Excellence: We apply careful legal research, disciplined preparation, and clear communication.",
        "Confidentiality: We protect client information and handle every matter with discretion.",
        "Service: We provide practical legal solutions that respect client time, budget, and objectives.",
        "Accountability: We maintain clear responsibilities, timely updates, and dependable follow-through.",
    ])
    add_heading(doc, "2.5 Areas of Specialization", 2)
    add_para(doc, "Vision Attorneys LF may customize this section to reflect the firm's approved practice areas, industry experience, and professional credentials.")
    add_bullet_list(doc, [
        "Corporate, commercial, and transactional advisory.",
        "Civil litigation, dispute resolution, and representation before courts and tribunals.",
        "Property, land, employment, family, and contract-related legal matters.",
        "Commissioner for Oaths services and legal documentation support.",
    ])


def add_services(doc):
    add_heading(doc, "3. Our Legal Services", 1)
    add_para(doc, "The following editable service descriptions can be adjusted to match the firm's actual scope of work, licensing, jurisdiction, and market positioning.")
    services = [
        ("CL", "Corporate and Commercial Law", "[Insert services related to company formation, governance, compliance, mergers, shareholder matters, and commercial transactions.]"),
        ("LT", "Civil Litigation", "[Insert litigation support for claims, defenses, pleadings, hearings, settlements, and enforcement proceedings.]"),
        ("CD", "Criminal Defense", "[Insert representation approach for suspects, accused persons, bail applications, trials, appeals, and related advisory work.]"),
        ("FL", "Family Law", "[Insert support for marriage, divorce, custody, maintenance, succession, adoption, and family settlement matters.]"),
        ("EL", "Employment and Labour Law", "[Insert advisory and representation services for employers, employees, contracts, disputes, termination, and compliance.]"),
        ("PL", "Property and Land Law", "[Insert land transactions, due diligence, leases, transfers, title review, disputes, and property advisory services.]"),
        ("CR", "Contract Drafting and Review", "[Insert contract preparation, review, negotiation support, risk clauses, and plain-language legal drafting.]"),
        ("AM", "Arbitration and Mediation", "[Insert alternative dispute resolution services, negotiation support, mediation representation, and arbitration preparation.]"),
        ("LA", "Legal Advisory Services", "[Insert advisory retainers, opinions, regulatory guidance, governance advice, and strategic legal risk management.]"),
        ("CO", "Commissioner for Oaths Services", "[Insert oath, affidavit, declaration, certification, and document witnessing services as permitted by law.]"),
        ("AS", "Additional Services", "[Insert any additional approved legal services, sector specializations, or emerging practice areas.]"),
    ]
    add_service_grid(doc, services)


def add_team(doc):
    add_heading(doc, "4. Our Team", 1)
    add_para(doc, "Replace the sample profiles below with actual biographies, professional photographs, qualifications, years of experience, bar admissions, and selected practice highlights.")
    add_team_profile(doc, "Managing Partner Profile Template", "[Insert Managing Partner Name]", "Managing Partner")
    add_para(doc, "", after=3)
    add_team_profile(doc, "Senior Advocate Profile Template", "[Insert Senior Advocate Name]", "Senior Advocate")
    add_para(doc, "", after=3)
    add_team_profile(doc, "Associate Advocate Profile Template", "[Insert Associate Advocate Name]", "Associate Advocate")
    add_para(doc, "", after=3)
    add_team_profile(doc, "Legal Assistant Profile Template", "[Insert Legal Assistant Name]", "Legal Assistant")


def add_why_choose(doc):
    add_heading(doc, "5. Why Choose Vision Attorneys LF", 1)
    rows = [
        ("Professionalism", "Every matter is handled with diligence, preparation, courtesy, and respect for legal procedure."),
        ("Integrity", "The firm is guided by ethical standards, honest advice, and responsible representation."),
        ("Client-Focused Approach", "Legal strategy is aligned with the client's objectives, risks, timeline, and resources."),
        ("Timely Legal Solutions", "The firm prioritizes clear communication, prompt action, and practical next steps."),
        ("Experienced Legal Team", "Clients benefit from multi-disciplinary legal knowledge and coordinated case handling."),
        ("Confidentiality and Trust", "Client information is protected with care, discretion, and professional responsibility."),
    ]
    add_key_value_table(doc, rows, widths=(2.05, 4.15), header_fill=LIGHT_GOLD)


def add_practice_areas(doc):
    add_heading(doc, "6. Key Practice Areas", 1)
    areas = [
        ("6.1 Commercial Law", "Vision Attorneys LF supports businesses and entrepreneurs with practical commercial legal guidance, including contract structures, governance documentation, regulatory obligations, transaction support, and dispute prevention. [Insert industry-specific capabilities and representative commercial matters.]"),
        ("6.2 Litigation", "The firm provides disciplined representation in civil, commercial, family, employment, and other disputes. Litigation services may include case assessment, pleadings, evidence preparation, court attendance, settlement negotiations, and enforcement support. [Insert court and tribunal experience.]"),
        ("6.3 Property Law", "Property law services may include land due diligence, sale agreements, leases, transfers, title review, boundary matters, landlord-tenant issues, and property dispute resolution. [Insert local land registry and property transaction experience.]"),
        ("6.4 Employment Law", "Vision Attorneys LF advises employers and employees on employment contracts, workplace policies, termination, disciplinary procedures, labour disputes, and compliance obligations. [Insert sector experience and dispute forum exposure.]"),
        ("6.5 Family Law", "The firm assists clients with sensitive family matters using a respectful, confidential, and solution-oriented approach. Services may include divorce, custody, maintenance, succession, adoption, and negotiated settlements. [Insert family-law credentials.]"),
    ]
    for heading, body in areas:
        add_heading(doc, heading, 2)
        add_para(doc, body)


def add_clients_cases_csr_partnerships(doc):
    add_heading(doc, "7. Representative Clients", 1)
    add_para(doc, "This section should be customized only with clients that the firm is authorized to identify publicly.")
    add_key_value_table(doc, [
        ("Government Institutions", "[Insert names or categories, subject to permission and confidentiality rules]"),
        ("Non-Governmental Organizations", "[Insert NGO clients or sectors served]"),
        ("Private Companies", "[Insert corporate clients, industries, or anonymized client categories]"),
        ("Individuals", "[Insert client groups served, such as entrepreneurs, families, employees, or landowners]"),
        ("International Organizations", "[Insert international bodies, development partners, or cross-border clients]"),
    ])
    add_heading(doc, "8. Selected Cases & Success Stories", 1)
    add_callout(doc, "Confidentiality Disclaimer", "Case information may be presented subject to client confidentiality obligations.")
    add_key_value_table(doc, [
        ("Case / Matter Title", "[Insert anonymized or approved case title]"),
        ("Client Sector", "[Insert sector or client category]"),
        ("Legal Issue", "[Insert brief description of the legal challenge]"),
        ("Firm Role", "[Insert advisory, litigation, negotiation, drafting, or compliance role]"),
        ("Outcome / Value", "[Insert result, settlement, risk avoided, or business value delivered]"),
    ])
    add_heading(doc, "9. Corporate Social Responsibility (CSR)", 1)
    add_bullet_list(doc, [
        "Community legal awareness programs: [Insert outreach initiatives, radio talks, workshops, or public education activities.]",
        "Pro bono services: [Insert pro bono policy, eligible client groups, and examples of public-interest support.]",
        "Youth mentorship initiatives: [Insert mentoring, internship, law student support, or career guidance programs.]",
    ])
    add_heading(doc, "10. Strategic Partnerships", 1)
    add_para(doc, "[Insert local and international partners, referral networks, professional associations, technical collaborators, and institutional relationships.]")
    add_heading(doc, "11. Firm Capacity and Competitive Advantage", 1)
    add_bullet_list(doc, [
        "Experienced legal professionals with complementary practice backgrounds.",
        "Multi-disciplinary expertise across commercial, litigation, property, employment, family, and advisory matters.",
        "Strong ethical standards and commitment to confidentiality.",
        "Efficient case management, clear client communication, and organized documentation.",
    ])


def add_contact_back_cover(doc):
    add_heading(doc, "12. Contact Information", 1)
    add_key_value_table(doc, [
        ("Physical Address", "[Insert Physical Address]"),
        ("Postal Address", "[Insert Postal Address]"),
        ("Telephone Numbers", "[Insert Telephone Numbers]"),
        ("Email Address", "[Insert Email Address]"),
        ("Website", "[Insert Website]"),
        ("Facebook", "[Insert Facebook Page URL]"),
        ("LinkedIn", "[Insert LinkedIn Page URL]"),
        ("Working Hours", "[Insert Working Hours]"),
    ], widths=(1.9, 4.3), header_fill=LIGHT_BLUE)
    doc.add_page_break()

    add_para(doc, "", after=28)
    add_placeholder_box(doc, "[Insert Firm Logo]", height_rows=1)
    add_para(doc, "Vision Attorneys LF", size=25, color=NAVY, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, before=18, after=2)
    add_para(doc, "(Advocates & Commissioner for Oaths)", size=12.5, color=GOLD, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=10)
    add_para(doc, "[Insert Motto]", size=13, color=MID_GRAY, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=22)
    add_callout(
        doc,
        "Professional Closing Statement",
        "Vision Attorneys LF welcomes the opportunity to provide reliable legal counsel, strategic representation, and client-centered service. [Insert final firm message or call to action.]",
    )
    add_para(doc, "", after=18)
    add_key_value_table(doc, [
        ("Address", "[Insert Physical Address]"),
        ("Phone", "[Insert Telephone Numbers]"),
        ("Email", "[Insert Email Address]"),
        ("Website", "[Insert Website]"),
    ], widths=(1.4, 4.5), header_fill=LIGHT_GOLD)
    add_para(doc, "", after=12)
    add_para(doc, "Prepared as an editable corporate profile template for Vision Attorneys LF.", size=9, color=MID_GRAY, align=WD_ALIGN_PARAGRAPH.CENTER)


def build_doc():
    doc = Document()
    configure_styles(doc)
    section = set_document_settings(doc)
    set_footer(section)
    add_cover(doc)
    add_toc(doc)
    add_executive_summary(doc)
    add_about(doc)
    add_services(doc)
    add_team(doc)
    add_why_choose(doc)
    add_practice_areas(doc)
    add_clients_cases_csr_partnerships(doc)
    doc.add_page_break()
    add_contact_back_cover(doc)
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build_doc())
